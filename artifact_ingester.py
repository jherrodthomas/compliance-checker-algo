"""
Universal Artifact Ingester for Compliance Engine

Ingests work products in multiple formats (.docx, .pdf, .md, .txt, .json) and
normalizes them into the standard artifact structure expected by the compliance engine.

Supports:
  - Single file ingestion
  - Directory scanning (all supported formats)
  - Format-specific text extraction
  - Smart auto-detection of clauses, tags, methods, and traceability
  - Robust error handling and encoding support
"""

import os
import sys
import json
import re
import subprocess
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Union


# Comprehensive automotive safety terminology for tag detection
SAFETY_TERMS = {
    # ISO 26262 concepts
    "asil", "safety goal", "safety requirement", "hazard", "hazard analysis",
    "hara", "fmea", "fmeda", "fta", "fault tree", "single point failure",
    "spf", "latent fault", "diagnostic coverage", "metric", "dcm", "dcmin",
    "mpf", "safe state", "safe system state", "residual risk",

    # Safety analysis methods
    "analysis", "verification", "validation", "testing", "review",
    "walkthrough", "inspection", "simulation", "model checking",
    "theorem proving", "formal verification", "code review",

    # E/E system concepts
    "hardware", "software", "middleware", "firmware", "architecture",
    "soc", "sub", "component", "interface", "requirement", "specification",
    "design", "implementation", "integration", "test", "qualification",

    # Safety lifecycle
    "concept", "planning", "development", "production", "operation",
    "maintenance", "decommissioning", "lifecycle", "phase", "stage",

    # Risk and failure concepts
    "risk", "severity", "probability", "occurrence", "failure",
    "failure mode", "failure rate", "mission time", "fault", "error",
    "defect", "anomaly", "degradation", "malfunction", "unavailability",

    # Automotive specific
    "vehicle", "ecu", "ecus", "sensor", "actuator", "control",
    "braking", "steering", "powertrain", "stability", "dynamics",
    "perception", "decision", "planning", "execution", "brake",
    "throttle", "transmission", "engine", "battery", "charging",

    # Traceability and documentation
    "requirement", "specification", "design", "implementation",
    "test case", "test plan", "test report", "traceability",
    "matrix", "mapping", "allocate", "allocation", "allocation document",

    # Quality and compliance
    "quality", "quality assurance", "qa", "process", "procedure",
    "metric", "measurement", "audit", "review", "assessment",
    "compliance", "standard", "norm", "regulation", "directive",

    # Configuration and change
    "configuration", "version", "release", "patch", "update",
    "change", "change management", "baseline", "milestone",

    # Other important terms
    "safety plan", "management", "organization", "responsibility",
    "competence", "training", "documentation", "evidence",
    "case", "dossier", "appraisal"
}

# Known safety analysis methods
SAFETY_METHODS = {
    "hara",                    # Hazard Analysis and Risk Assessment
    "hazard analysis",
    "hazard analysis and risk assessment",
    "fmea",                    # Failure Modes and Effects Analysis
    "failure modes and effects analysis",
    "fmeda",                   # Failure Modes, Effects and Diagnostics Analysis
    "failure modes and effects diagnostics analysis",
    "fta",                     # Fault Tree Analysis
    "fault tree analysis",
    "fault tree",
    "hazop",                   # Hazard and Operability
    "hazard and operability",
    "dfa",                     # Design Failure Analysis
    "design failure analysis",
    "fea",                     # Finite Element Analysis
    "finite element analysis",
    "cea",                     # Common Cause Analysis
    "common cause analysis",
    "sneak circuit analysis",
    "sneak",
    "emi analysis",
    "electromagnetic interference",
    "thermal analysis",
    "shock and vibration analysis",
    "stress analysis",
    "code review",
    "static analysis",
    "dynamic testing",
    "unit testing",
    "integration testing",
    "system testing",
    "acceptance testing",
    "regression testing",
    "simulation",
    "model checking",
    "theorem proving",
    "formal verification",
    "walkthrough",
    "inspection",
    "technical review",
    "management review",
    "appraisal",
    "assessment"
}


class UniversalArtifactIngester:
    """
    Universal ingester for work products in multiple formats.

    Handles .docx, .pdf, .md, .txt, and .json files, extracting content
    and normalizing them into the compliance engine's standard artifact structure.
    """

    SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.md', '.txt', '.json', '.xlsx', '.xls', '.csv', '.odt'}

    def __init__(self, verbose: bool = False):
        """
        Initialize the ingester.

        Args:
            verbose: If True, print status messages during ingestion
        """
        self.verbose = verbose
        self.work_products = []

    def ingest(self, path: str) -> Dict:
        """
        Ingest a file or directory of files.

        If path is a file, ingests that single file.
        If path is a directory, scans for all supported file types and ingests them.

        Args:
            path: File or directory path

        Returns:
            Dict with structure: {"work_products": [list of normalized artifacts]}
        """
        self.work_products = []
        path_obj = Path(path)

        if not path_obj.exists():
            raise FileNotFoundError(f"Path does not exist: {path}")

        files_to_ingest = []

        if path_obj.is_file():
            files_to_ingest = [path_obj]
        else:
            # Scan directory for supported files
            for ext in self.SUPPORTED_EXTENSIONS:
                files_to_ingest.extend(sorted(path_obj.glob(f"*{ext}")))
                files_to_ingest.extend(sorted(path_obj.glob(f"**/*{ext}")))

            # Remove duplicates while preserving order
            seen = set()
            unique_files = []
            for f in files_to_ingest:
                if f not in seen:
                    seen.add(f)
                    unique_files.append(f)
            files_to_ingest = unique_files

        if not files_to_ingest:
            if self.verbose:
                print(f"No supported files found in {path}")
            return {"work_products": []}

        if self.verbose:
            print(f"Found {len(files_to_ingest)} file(s) to ingest")

        # Ingest each file
        for i, filepath in enumerate(files_to_ingest, 1):
            try:
                wp_id = f"WP-{i:03d}"
                artifact = self.ingest_file(str(filepath), wp_id)
                if artifact:
                    self.work_products.append(artifact)
                    if self.verbose:
                        print(f"  [{wp_id}] {filepath.name}")
            except Exception as e:
                if self.verbose:
                    print(f"  ERROR: {filepath.name} - {str(e)}")
                continue

        return {"work_products": self.work_products}

    def ingest_file(self, filepath: str, wp_id: str = "WP-001") -> Optional[Dict]:
        """
        Ingest a single file of any supported format.

        Args:
            filepath: Path to the file
            wp_id: Work product ID (e.g., "WP-001")

        Returns:
            Normalized work product dict, or None if ingestion fails
        """
        filepath = str(filepath)
        path_obj = Path(filepath)

        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        ext = path_obj.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}")

        # Extract content based on format
        if ext == '.docx':
            title, content = self._extract_docx(filepath)
        elif ext == '.pdf':
            title, content = self._extract_pdf(filepath)
        elif ext == '.md':
            title, content = self._extract_markdown(filepath)
        elif ext == '.txt':
            title, content = self._extract_text(filepath)
        elif ext == '.json':
            title, content = self._extract_json(filepath)
        elif ext in ('.xlsx', '.xls', '.csv'):
            title, content = self._extract_spreadsheet(filepath)
        elif ext == '.odt':
            title, content = self._extract_odt(filepath)
        else:
            return None

        # Normalize the extracted content
        artifact = self._normalize(wp_id, title, content, filepath)
        return artifact

    def _extract_docx(self, filepath: str) -> Tuple[str, str]:
        """
        Extract title and content from a .docx file.

        Args:
            filepath: Path to the .docx file

        Returns:
            Tuple of (title, content)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for .docx support")

        doc = Document(filepath)

        # Try to get title from document properties
        title = doc.core_properties.title

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract table contents
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    paragraphs.append(" | ".join(row_cells))

        content = "\n".join(paragraphs)

        # If no title in properties, use first heading or first line
        if not title:
            if paragraphs:
                title = paragraphs[0][:100]
            else:
                title = Path(filepath).stem

        return title, content

    def _extract_pdf(self, filepath: str) -> Tuple[str, str]:
        """
        Extract title and content from a .pdf file using pdftotext.

        Args:
            filepath: Path to the .pdf file

        Returns:
            Tuple of (title, content)
        """
        try:
            # Try with -layout first for better text extraction
            result = subprocess.run(
                ['pdftotext', '-layout', filepath, '-'],
                capture_output=True,
                text=True,
                timeout=30
            )

            if result.returncode != 0:
                # Fallback without layout
                result = subprocess.run(
                    ['pdftotext', filepath, '-'],
                    capture_output=True,
                    text=True,
                    timeout=30
                )

            content = result.stdout

            # If extracted text is very short, try alternative approach
            if len(content.strip()) < 500:
                result = subprocess.run(
                    ['pdftotext', filepath, '-'],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                content = result.stdout

        except FileNotFoundError:
            raise RuntimeError("pdftotext command not found. Install poppler-utils.")
        except subprocess.TimeoutExpired:
            raise RuntimeError(f"pdftotext timed out processing {filepath}")

        # Use filename as title
        title = Path(filepath).stem

        return title, content

    def _extract_markdown(self, filepath: str) -> Tuple[str, str]:
        """
        Extract title and content from a .md file.

        Args:
            filepath: Path to the .md file

        Returns:
            Tuple of (title, content)
        """
        content = self._read_file(filepath)

        # Try to extract title from first h1 heading
        title = None
        h1_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if h1_match:
            title = h1_match.group(1)

        if not title:
            title = Path(filepath).stem

        return title, content

    def _extract_text(self, filepath: str) -> Tuple[str, str]:
        """
        Extract title and content from a .txt file.

        Args:
            filepath: Path to the .txt file

        Returns:
            Tuple of (title, content)
        """
        content = self._read_file(filepath)

        # Try to detect title from first line or section heading
        lines = content.split('\n')
        title = None

        # Look for ALL CAPS lines or numbered section at start
        for line in lines[:5]:
            stripped = line.strip()
            if stripped and (stripped.isupper() or re.match(r'^\d+\.\s+', stripped)):
                title = stripped[:100]
                break

        if not title:
            # Use first non-empty line
            for line in lines:
                if line.strip():
                    title = line.strip()[:100]
                    break

        if not title:
            title = Path(filepath).stem

        return title, content

    def _extract_json(self, filepath: str) -> Tuple[str, str]:
        """
        Extract title and content from a .json file.

        If the file matches the work_products schema, treat it as pre-normalized.
        Otherwise, treat field values as content.

        Args:
            filepath: Path to the .json file

        Returns:
            Tuple of (title, content)
        """
        content = self._read_file(filepath)

        try:
            data = json.loads(content)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON: {e}")

        # Check if this is already in the expected work_products format
        if isinstance(data, dict) and "work_products" in data:
            # Return as-is, will be handled specially
            return "Pre-normalized JSON", json.dumps(data, indent=2)

        # Otherwise, extract all text content from JSON values
        text_parts = []
        title = Path(filepath).stem

        def extract_values(obj):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if isinstance(value, str) and value.strip():
                        text_parts.append(value)
                    else:
                        extract_values(value)
            elif isinstance(obj, list):
                for item in obj:
                    extract_values(item)
            elif isinstance(obj, str) and obj.strip():
                text_parts.append(obj)

        extract_values(data)
        extracted_content = "\n".join(text_parts)

        return title, extracted_content

    def _extract_odt(self, filepath: str) -> Tuple[str, str]:
        """
        Extract title and content from an .odt (OpenDocument Text) file.

        ODT files are ZIP archives containing content.xml with the document text.
        Uses only Python stdlib (zipfile + xml.etree).

        Args:
            filepath: Path to the .odt file

        Returns:
            Tuple of (title, content)
        """
        import zipfile
        from xml.etree import ElementTree as ET

        title = Path(filepath).stem

        with zipfile.ZipFile(filepath, 'r') as z:
            # Read content.xml
            with z.open('content.xml') as f:
                tree = ET.parse(f)

        root = tree.getroot()

        # ODF namespaces
        ns = {
            'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
            'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
        }

        # Extract all text content
        paragraphs = []
        for elem in root.iter():
            # Get paragraphs and headings
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag in ('p', 'h'):
                # Gather all text in this element (including child spans)
                text = ''.join(elem.itertext()).strip()
                if text:
                    paragraphs.append(text)

        content = '\n'.join(paragraphs)

        # Try to detect title from first heading or paragraph
        if paragraphs:
            first = paragraphs[0]
            if len(first) < 200:
                title = first

        return title, content

    def _extract_spreadsheet(self, filepath: str) -> Tuple[str, str]:
        """
        Extract title and content from spreadsheet files (.xlsx, .xls, .csv).

        For .xlsx/.xls: uses openpyxl (preferred) or falls back to csv via
        subprocess (ssconvert from gnumeric, or xlrd).
        For .csv: reads directly.

        Each row is converted to a text line. Headers become field labels.

        Args:
            filepath: Path to the spreadsheet file

        Returns:
            Tuple of (title, content)
        """
        ext = Path(filepath).suffix.lower()
        title = Path(filepath).stem

        if ext == '.csv':
            import csv
            rows = []
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(row)
            return title, self._rows_to_text(rows)

        # .xlsx / .xls — try openpyxl first
        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath, read_only=True, data_only=True)
            all_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                all_text.append(f"=== Sheet: {sheet_name} ===")
                rows = []
                for row in ws.iter_rows(values_only=True):
                    rows.append([str(cell) if cell is not None else "" for cell in row])
                all_text.append(self._rows_to_text(rows))
            wb.close()
            return title, "\n\n".join(all_text)
        except ImportError:
            pass

        # Fallback: try ssconvert (gnumeric) to CSV, then parse
        if subprocess.run(["which", "ssconvert"], capture_output=True).returncode == 0:
            import tempfile, csv
            tmp_csv = tempfile.mktemp(suffix=".csv")
            try:
                subprocess.run(["ssconvert", filepath, tmp_csv], check=True, capture_output=True)
                rows = []
                with open(tmp_csv, 'r', encoding='utf-8', errors='replace') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        rows.append(row)
                return title, self._rows_to_text(rows)
            finally:
                if os.path.exists(tmp_csv):
                    os.remove(tmp_csv)

        # Last resort: read raw bytes for any text content
        content = self._read_file(filepath)
        if not content.strip():
            raise ImportError(
                "Cannot read .xlsx files. Install openpyxl:\n"
                "  pip install openpyxl"
            )
        return title, content

    @staticmethod
    def _rows_to_text(rows: List[List[str]]) -> str:
        """
        Convert spreadsheet rows to readable text.

        If headers are detected (first row), formats as 'Header: Value' pairs.
        Otherwise formats as tab-separated lines.

        Args:
            rows: List of row lists

        Returns:
            Formatted text string
        """
        if not rows:
            return ""

        # Check if first row looks like headers (all non-empty, all short text)
        headers = rows[0] if rows else []
        is_header = (
            len(headers) >= 2 and
            all(h.strip() for h in headers) and
            all(len(h) < 100 for h in headers)
        )

        lines = []
        if is_header and len(rows) > 1:
            for row in rows[1:]:
                parts = []
                for i, cell in enumerate(row):
                    if cell.strip():
                        label = headers[i] if i < len(headers) else f"Col{i+1}"
                        parts.append(f"{label}: {cell}")
                if parts:
                    lines.append(" | ".join(parts))
        else:
            for row in rows:
                line = "\t".join(cell for cell in row if cell.strip())
                if line.strip():
                    lines.append(line)

        return "\n".join(lines)

    def _normalize(self, wp_id: str, title: str, content: str, filepath: str) -> Dict:
        """
        Normalize extracted content into standard work product structure.

        Args:
            wp_id: Work product ID
            title: Document title
            content: Full text content
            filepath: Original file path (for reference)

        Returns:
            Normalized work product dict
        """
        return {
            "id": wp_id,
            "title": title,
            "status": "complete",
            "content": content,
            "mapped_clauses": self._detect_clauses(content),
            "tags": self._detect_tags(content),
            "methods_used": self._detect_methods(content),
            "traceability": self._detect_traceability(content),
            "source_file": filepath
        }

    def _detect_clauses(self, text: str) -> List[str]:
        """
        Detect ISO clause references in text.

        Looks for patterns like:
          - "6.4.2", "clause 6.4.2", "per 6.4.2"
          - "ISO 26262-3 7.4.1"
          - "IEC 61508-3 Section 4.2"

        Args:
            text: Content to search

        Returns:
            List of detected clause references
        """
        clauses = set()

        # Pattern 1: ISO/IEC standard with section (e.g., "ISO 26262-3 7.4.1")
        pattern1 = r'(?:ISO|IEC)\s+\d+(?:-\d+)?\s+([0-9]+(?:\.[0-9]+)*)'
        matches1 = re.finditer(pattern1, text, re.IGNORECASE)
        for match in matches1:
            clauses.add(match.group(1))

        # Pattern 2: clause/section reference (e.g., "clause 5.4.2" or "per 5.4.2")
        pattern2 = r'(?:clause|section|per|subsection|sub-section)\s+([0-9]+(?:\.[0-9]+)*)'
        matches2 = re.finditer(pattern2, text, re.IGNORECASE)
        for match in matches2:
            clauses.add(match.group(1))

        # Pattern 3: standalone clause numbers (e.g., "5.4.2" when preceded by keywords)
        pattern3 = r'(?:requires?|satisfies?|per|derived?\s+from|traced?\s+to|mapped?\s+to)\s+([0-9]+(?:\.[0-9]+)*)'
        matches3 = re.finditer(pattern3, text, re.IGNORECASE)
        for match in matches3:
            clauses.add(match.group(1))

        return sorted(list(clauses))

    def _detect_tags(self, text: str) -> List[str]:
        """
        Extract safety/automotive concept tags from text.

        Looks for known safety terminology and returns unique, sorted tags.

        Args:
            text: Content to search

        Returns:
            List of detected tags
        """
        tags = set()
        text_lower = text.lower()

        for term in SAFETY_TERMS:
            # Use word boundaries to avoid partial matches
            pattern = r'\b' + re.escape(term) + r'\b'
            if re.search(pattern, text_lower):
                tags.add(term)

        return sorted(list(tags))

    def _detect_methods(self, text: str) -> List[str]:
        """
        Find known safety analysis methods mentioned in text.

        Args:
            text: Content to search

        Returns:
            List of detected methods
        """
        methods = set()
        text_lower = text.lower()

        for method in SAFETY_METHODS:
            # Use word boundaries
            pattern = r'\b' + re.escape(method) + r'\b'
            if re.search(pattern, text_lower):
                methods.add(method)

        return sorted(list(methods))

    def _detect_traceability(self, text: str) -> Dict:
        """
        Detect traceability relationships in text.

        Looks for keywords like:
          - "derives from", "derived from"
          - "traced to", "traces to"
          - "satisfies", "satisfied by"
          - "verified by", "verifies"
          - "implemented by", "implements"

        Args:
            text: Content to search

        Returns:
            Dict with 'from' and 'to' keys containing detected relationships
        """
        traceability = {"from": [], "to": []}

        # Pattern: "derives from X" or "derived from X"
        derives_pattern = r'(?:derives?|derived)\s+from\s+([A-Za-z0-9\s\-_]+?)(?:\.|,|;|$|\n)'
        for match in re.finditer(derives_pattern, text, re.IGNORECASE):
            source = match.group(1).strip()
            if source and len(source) < 200:
                if source not in traceability["from"]:
                    traceability["from"].append(source)

        # Pattern: "traced to X" or "traces to X"
        traces_pattern = r'(?:traced?|traces)\s+to\s+([A-Za-z0-9\s\-_]+?)(?:\.|,|;|$|\n)'
        for match in re.finditer(traces_pattern, text, re.IGNORECASE):
            target = match.group(1).strip()
            if target and len(target) < 200:
                if target not in traceability["to"]:
                    traceability["to"].append(target)

        # Pattern: "satisfies X"
        satisfies_pattern = r'satisfies?\s+([A-Za-z0-9\s\-_]+?)(?:\.|,|;|$|\n)'
        for match in re.finditer(satisfies_pattern, text, re.IGNORECASE):
            target = match.group(1).strip()
            if target and len(target) < 200:
                if target not in traceability["to"]:
                    traceability["to"].append(target)

        # Pattern: "verified by X"
        verified_pattern = r'verified?\s+by\s+([A-Za-z0-9\s\-_]+?)(?:\.|,|;|$|\n)'
        for match in re.finditer(verified_pattern, text, re.IGNORECASE):
            target = match.group(1).strip()
            if target and len(target) < 200:
                if target not in traceability["to"]:
                    traceability["to"].append(target)

        return traceability

    def _read_file(self, filepath: str, encoding: str = 'utf-8') -> str:
        """
        Read file with graceful encoding fallback.

        Tries UTF-8 first, falls back to latin-1.

        Args:
            filepath: Path to file
            encoding: Initial encoding to try

        Returns:
            File contents as string
        """
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 (should never fail)
            with open(filepath, 'r', encoding='latin-1') as f:
                return f.read()


def main():
    """
    CLI interface for the artifact ingester.

    Usage:
        python artifact_ingester.py <file_or_dir> [output.json]
    """
    if len(sys.argv) < 2:
        print("Usage: python artifact_ingester.py <file_or_dir> [output.json]")
        print("\nSupported formats: .docx, .pdf, .md, .txt, .json")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    ingester = UniversalArtifactIngester(verbose=True)

    try:
        result = ingester.ingest(input_path)

        print(f"\nSuccessfully ingested {len(result['work_products'])} artifact(s)")

        if output_path:
            with open(output_path, 'w') as f:
                json.dump(result, f, indent=2)
            print(f"Output written to: {output_path}")
        else:
            print("\nNormalized artifacts (JSON):")
            print(json.dumps(result, indent=2))

    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
