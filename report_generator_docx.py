#!/usr/bin/env python3
"""
ComplianceIQ — Professional Compliance Assessment Report Generator (DOCX)
=========================================================================
Generates audit-grade compliance assessment reports modelled on TUV/DEKRA
functional safety assessment (FSA) documents and scientific journal standards.

Report Structure:
  Cover Page · Document Control · Table of Contents
  1. Executive Summary
  2. Assessment Scope & Methodology
  3. Work Product Register
  4. Compliance Assessment Dashboard
  5. Requirements Compliance Matrix
  6. Traceability Matrix
  7. Gap Analysis
  8. Method & Practice Audit
  9. Risk Assessment
  10. Corrective Action Register
  11. Confirmation Review (ISO 26262-2:2018 §6.4.4)
  12. Verification Review (ISO 26262-2:2018 §6.4.10 / Part 8)
  13. Assessment Decision
  Appendix A: Detailed Findings
  Appendix B: Algorithm Layer Details
  Appendix C: Glossary & References

Usage:
  python report_generator_docx.py compliance_report_agnostic.json [output.docx]

Or from code:
  from report_generator_docx import generate_compliance_docx
  generate_compliance_docx(report_dict, "output.docx")
"""

import json
import os
import sys
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Optional: matplotlib for high-quality charts
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.ticker as mticker
    import matplotlib.patches as mpatches
    from matplotlib.path import Path
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

# Optional: seaborn for enhanced statistical visualizations
try:
    import seaborn as sns
    HAS_SEABORN = True
except ImportError:
    HAS_SEABORN = False

# Optional: numpy for numerical operations
try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False

# Optional: graphviz for tree diagrams
try:
    import graphviz
    HAS_GRAPHVIZ = True
except ImportError:
    HAS_GRAPHVIZ = False

# ═══════════════════════════════════════════════════
#  BRANDING
# ═══════════════════════════════════════════════════

BRAND_NAME = "Lion of Functional Safety Engine\u2122"
BRAND_SHORT = "ComplianceIQ"
BRAND_TAGLINE = "Lion of Functional Safety Engine"

# Logo search paths (checked in order)
LOGO_SEARCH_PATHS = [
    # User's specified path
    r"C:\Users\Jrod\Downloads\The Lion V2_\The Lion V2\TL-V2-White-BG.png",
    # Common relative paths
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png"),
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "TL-V2-White-BG.png"),
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "logo.png"),
]


# ═══════════════════════════════════════════════════
#  COLOR PALETTE — Professional audit document
# ═══════════════════════════════════════════════════

NAVY       = RGBColor(0x0D, 0x2B, 0x4E)
DARK_BLUE  = RGBColor(0x1A, 0x47, 0x80)
MID_BLUE   = RGBColor(0x2E, 0x6E, 0xB5)
LIGHT_BLUE = RGBColor(0xD9, 0xE6, 0xF2)
ACCENT     = RGBColor(0xD4, 0x6B, 0x08)  # Warm amber
GREEN_OK   = RGBColor(0x16, 0x6B, 0x34)
YELLOW_WARN = RGBColor(0xCA, 0x8A, 0x04)
RED_CRIT   = RGBColor(0xB9, 0x1C, 0x1C)
GREY_TEXT  = RGBColor(0x4B, 0x55, 0x63)
GREY_LIGHT = RGBColor(0x9C, 0xA3, 0xAF)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x1F, 0x2A, 0x37)
BG_ALT     = "F3F4F6"   # Alternating row background
HDR_BG     = "0D2B4E"   # Header background
HDR_BG2    = "1A4780"   # Secondary header


# ═══════════════════════════════════════════════════
#  HELPERS
# ═══════════════════════════════════════════════════

def score_color(score):
    if score >= 80: return GREEN_OK
    if score >= 60: return YELLOW_WARN
    return RED_CRIT

def verdict_color(verdict):
    v = verdict.upper()
    if "NOT COMPLIANT" in v and "CONDITIONALLY" not in v:
        return RED_CRIT
    if "NOT ACCEPTABLE" in v and "CONDITIONALLY" not in v:
        return RED_CRIT
    if "CONDITIONALLY" in v:
        return ACCENT
    return GREEN_OK

def severity_color(sev):
    sev = (sev or "").upper()
    if sev == "CRITICAL": return RED_CRIT
    if sev == "MAJOR": return ACCENT
    if sev == "WARNING": return YELLOW_WARN
    return GREY_LIGHT

def status_color(status):
    s = (status or "").lower()
    if "not compliant" in s or "non-compliant" in s:
        return RED_CRIT
    if "missing" in s:
        return RED_CRIT
    if "partial" in s or "conditional" in s:
        return YELLOW_WARN
    if s == "n/a":
        return GREY_LIGHT
    if "compliant" in s or "acceptable" in s:
        return GREEN_OK
    return GREY_TEXT

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        sz, style, color = val.split()
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=0, bottom=0, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{left}" w:type="dxa"/>'
        f'<w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)

def add_run(paragraph, text, bold=False, size=None, color=None, italic=False, font_name=None):
    run = paragraph.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if font_name: run.font.name = font_name
    return run

def add_heading(doc, text, level=1, color=NAVY):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
    return heading

def add_hr(doc, color_hex="0D2B4E", weight="6"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{weight}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_thin_hr(doc):
    return add_hr(doc, color_hex="D1D5DB", weight="4")

def clear_borders(cell):
    set_cell_border(cell, top="0 none FFFFFF", bottom="0 none FFFFFF",
                    left="0 none FFFFFF", right="0 none FFFFFF")

def style_header_row(row, bg=HDR_BG):
    for cell in row.cells:
        set_cell_shading(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def style_data_row(row, idx):
    if idx % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, BG_ALT)

def find_logo():
    """Search for the brand logo in known locations."""
    for path in LOGO_SEARCH_PATHS:
        if os.path.isfile(path):
            return path
    return None


# ═══════════════════════════════════════════════════
#  CHART GENERATION (matplotlib — publication-quality)
# ═══════════════════════════════════════════════════

CHART_DPI = 300  # Publication-quality resolution
CHART_BG = '#FAFBFC'

def _chart_colors():
    return {
        'navy': '#0D2B4E', 'blue': '#2E6EB5', 'green': '#166B34',
        'yellow': '#CA8A04', 'red': '#B91C1C', 'amber': '#D46B08',
        'grey': '#9CA3AF', 'bg': CHART_BG, 'light_grey': '#E5E7EB',
        'dark_text': '#1F2A37',
    }

def _apply_professional_style():
    """Apply a consistent professional style to all charts."""
    if not HAS_MATPLOTLIB:
        return
    plt.rcParams.update({
        'font.family': 'sans-serif',
        'font.sans-serif': ['DejaVu Sans', 'Helvetica', 'Arial', 'sans-serif'],
        'font.size': 9,
        'axes.titlesize': 11,
        'axes.labelsize': 9,
        'xtick.labelsize': 8,
        'ytick.labelsize': 8,
        'legend.fontsize': 8,
        'figure.dpi': CHART_DPI,
        'savefig.dpi': CHART_DPI,
        'axes.spines.top': False,
        'axes.spines.right': False,
        'axes.edgecolor': '#D1D5DB',
        'axes.grid': True,
        'grid.alpha': 0.25,
        'grid.color': '#D1D5DB',
        'grid.linestyle': '--',
    })


def generate_layer_chart(layers_data, layer_labels, layer_weights):
    """Generate a horizontal bar chart of algorithm layer scores. Returns temp PNG path."""
    if not HAS_MATPLOTLIB or not layers_data:
        return None
    try:
        _apply_professional_style()
        colors = _chart_colors()
        labels = []
        scores = []
        weights = []
        bar_colors = []
        for key in reversed(list(layer_labels.keys())):
            if key in layers_data:
                lbl = layer_labels[key]
                sc = layers_data[key].get("score", 0)
                w = layer_weights.get(key, 0)
                labels.append(lbl)
                scores.append(sc)
                weights.append(w)
                if sc >= 80: bar_colors.append(colors['green'])
                elif sc >= 60: bar_colors.append(colors['yellow'])
                else: bar_colors.append(colors['red'])

        fig, ax = plt.subplots(figsize=(6.5, 3.2))
        fig.patch.set_facecolor(colors['bg'])
        ax.set_facecolor(colors['bg'])

        # Gradient-style bars with rounded edges
        bars = ax.barh(labels, scores, color=bar_colors, height=0.55,
                       edgecolor='white', linewidth=0.8, alpha=0.9,
                       zorder=3)

        # Background reference bars at 100%
        ax.barh(labels, [100]*len(labels), color=colors['light_grey'],
                height=0.55, alpha=0.35, zorder=1)

        # Score labels on bars
        for bar, sc, w in zip(bars, scores, weights):
            x_pos = bar.get_width() + 1.5
            if x_pos > 95:
                x_pos = bar.get_width() - 8
                clr = 'white'
            else:
                clr = colors['navy']
            ax.text(x_pos, bar.get_y() + bar.get_height()/2,
                    f'{sc:.0f}%', va='center', ha='left', fontsize=8.5,
                    fontweight='bold', color=clr, zorder=5)
            # Weight annotation
            ax.text(102, bar.get_y() + bar.get_height()/2,
                    f'×{w:.0%}', va='center', ha='left', fontsize=6.5,
                    color=colors['grey'], style='italic', zorder=5)

        ax.set_xlim(0, 112)
        ax.set_xlabel('Score (%)', fontsize=9, color=colors['navy'], fontweight='bold')
        ax.tick_params(axis='y', labelsize=8.5, colors=colors['navy'])
        ax.tick_params(axis='x', labelsize=7.5, colors=colors['grey'])
        ax.xaxis.set_major_locator(mticker.MultipleLocator(25))
        ax.spines['left'].set_color(colors['light_grey'])
        ax.spines['bottom'].set_color(colors['light_grey'])

        # Title
        ax.set_title('Algorithm Layer Performance', fontsize=11, fontweight='bold',
                     color=colors['navy'], pad=12)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_layer_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Layer chart generation failed: {e}")
        return None


def generate_severity_pie(sev_counts):
    """Generate a severity distribution donut chart. Returns temp PNG path."""
    if not HAS_MATPLOTLIB:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()
        items = [
            ("Critical", sev_counts.get("CRITICAL", 0), colors_map['red']),
            ("Major", sev_counts.get("MAJOR", 0), colors_map['amber']),
            ("Warning", sev_counts.get("WARNING", 0), colors_map['yellow']),
            ("Info", sev_counts.get("INFO", 0), colors_map['grey']),
        ]
        items = [(l, v, c) for l, v, c in items if v > 0]
        if not items:
            return None

        labels, values, pie_colors = zip(*items)

        fig, ax = plt.subplots(figsize=(3.5, 3.5))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])

        wedges, texts, autotexts = ax.pie(
            values, labels=None, colors=pie_colors, autopct='%1.0f%%',
            startangle=90, pctdistance=0.78,
            wedgeprops=dict(width=0.38, edgecolor='white', linewidth=2.5),
            shadow=False)

        for t in autotexts:
            t.set_fontsize(8.5)
            t.set_fontweight('bold')
            t.set_color('white')

        # Center text — total findings
        total = sum(values)
        ax.text(0, 0.05, f'{total}', ha='center', va='center',
                fontsize=22, fontweight='bold', color=colors_map['navy'])
        ax.text(0, -0.14, 'findings', ha='center', va='center',
                fontsize=8.5, color=colors_map['grey'])

        ax.legend(labels, loc='lower center', ncol=min(len(labels), 4),
                  fontsize=7.5, frameon=False,
                  bbox_to_anchor=(0.5, -0.06))

        ax.set_title('Finding Severity Distribution', fontsize=10, fontweight='bold',
                     color=colors_map['navy'], pad=8)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_sev_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Severity chart generation failed: {e}")
        return None


def generate_compliance_gauge(score):
    """Generate a semi-circular gauge chart for compliance score. Returns temp PNG path."""
    if not HAS_MATPLOTLIB:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()
        import numpy as np

        fig, ax = plt.subplots(figsize=(4.0, 2.5))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])
        ax.set_aspect('equal')

        if score >= 80: fill_color = colors_map['green']
        elif score >= 60: fill_color = colors_map['yellow']
        else: fill_color = colors_map['red']

        # Draw gauge arcs
        theta_bg = np.linspace(0, np.pi, 100)
        r_outer, r_inner = 1.0, 0.65
        # Background arc
        x_bg = np.concatenate([r_outer * np.cos(theta_bg), r_inner * np.cos(theta_bg[::-1])])
        y_bg = np.concatenate([r_outer * np.sin(theta_bg), r_inner * np.sin(theta_bg[::-1])])
        ax.fill(x_bg, y_bg, color=colors_map['light_grey'], alpha=0.5)

        # Filled arc (score)
        theta_fill = np.linspace(np.pi, np.pi - (score/100)*np.pi, 100)
        x_fill = np.concatenate([r_outer * np.cos(theta_fill), r_inner * np.cos(theta_fill[::-1])])
        y_fill = np.concatenate([r_outer * np.sin(theta_fill), r_inner * np.sin(theta_fill[::-1])])
        ax.fill(x_fill, y_fill, color=fill_color, alpha=0.9)

        # Score text
        ax.text(0, 0.35, f'{score:.0f}%', ha='center', va='center',
                fontsize=28, fontweight='bold', color=colors_map['navy'])
        ax.text(0, 0.05, 'Compliance Score', ha='center', va='center',
                fontsize=9, color=colors_map['grey'])

        # Grade labels
        ax.text(-1.05, -0.05, '0%', ha='center', fontsize=7, color=colors_map['grey'])
        ax.text(1.05, -0.05, '100%', ha='center', fontsize=7, color=colors_map['grey'])

        ax.set_xlim(-1.3, 1.3)
        ax.set_ylim(-0.15, 1.15)
        ax.axis('off')

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_gauge_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Gauge chart generation failed: {e}")
        return None


def generate_trace_coverage_chart(trace_cov):
    """Generate a traceability coverage bar chart. Returns temp PNG path."""
    if not HAS_MATPLOTLIB:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()
        total = trace_cov.get("total_requirements", 0)
        traced = trace_cov.get("traced_requirements", 0)
        orphan_r = trace_cov.get("orphan_requirements_count", 0)
        orphan_a = trace_cov.get("orphan_artifacts_count", 0)

        if total == 0 and traced == 0:
            return None

        cats = ['Requirements\nTraced', 'Orphan\nRequirements', 'Orphan\nArtifacts']
        vals = [traced, orphan_r, orphan_a]
        bar_colors = [colors_map['green'], colors_map['red'], colors_map['amber']]

        fig, ax = plt.subplots(figsize=(5.0, 2.8))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])

        bars = ax.bar(cats, vals, color=bar_colors, width=0.5,
                      edgecolor='white', linewidth=1.0, alpha=0.9, zorder=3)
        for bar, val in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                    str(val), ha='center', va='bottom', fontsize=10,
                    fontweight='bold', color=colors_map['navy'], zorder=5)

        ax.set_ylabel('Count', fontsize=9, color=colors_map['navy'], fontweight='bold')
        ax.tick_params(axis='both', labelsize=8, colors=colors_map['grey'])
        ax.spines['left'].set_color(colors_map['light_grey'])
        ax.spines['bottom'].set_color(colors_map['light_grey'])
        ax.set_title('Traceability Coverage', fontsize=11, fontweight='bold',
                     color=colors_map['navy'], pad=12)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_trace_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Trace chart generation failed: {e}")
        return None


def generate_confirmation_review_chart(conf_summary):
    """Generate a confirmation review results chart. Returns temp PNG path."""
    if not HAS_MATPLOTLIB:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()

        passed = conf_summary.get("pass", 0)
        partial = conf_summary.get("partial", 0)
        failed = conf_summary.get("fail", 0)

        if passed + partial + failed == 0:
            return None

        cats = ['Pass', 'Partial', 'Fail']
        vals = [passed, partial, failed]
        bar_colors = [colors_map['green'], colors_map['yellow'], colors_map['red']]

        fig, ax = plt.subplots(figsize=(4.0, 2.5))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])

        bars = ax.bar(cats, vals, color=bar_colors, width=0.5,
                      edgecolor='white', linewidth=1.0, alpha=0.9, zorder=3)
        for bar, val in zip(bars, vals):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.15,
                        str(val), ha='center', va='bottom', fontsize=11,
                        fontweight='bold', color=colors_map['navy'], zorder=5)

        ax.set_ylabel('Checklist Items', fontsize=9, color=colors_map['navy'], fontweight='bold')
        ax.tick_params(axis='both', labelsize=8, colors=colors_map['grey'])
        ax.set_title('Confirmation Review Results', fontsize=11, fontweight='bold',
                     color=colors_map['navy'], pad=12)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_conf_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Confirmation review chart generation failed: {e}")
        return None


def generate_verification_coverage_chart(verif_wps):
    """Generate a verification coverage chart per work product. Returns temp PNG path."""
    if not HAS_MATPLOTLIB or not verif_wps:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()

        # Show up to 12 work products
        wps = verif_wps[:12]
        labels = [w.get("name", "WP")[:40] for w in wps]
        coverages = [w.get("coverage_pct", 0) for w in wps]
        bar_colors = []
        for c in coverages:
            if c >= 75: bar_colors.append(colors_map['green'])
            elif c >= 40: bar_colors.append(colors_map['yellow'])
            else: bar_colors.append(colors_map['red'])

        fig, ax = plt.subplots(figsize=(6.5, max(2.5, len(wps) * 0.35)))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])

        bars = ax.barh(labels, coverages, color=bar_colors, height=0.5,
                       edgecolor='white', linewidth=0.8, alpha=0.9, zorder=3)
        ax.barh(labels, [100]*len(labels), color=colors_map['light_grey'],
                height=0.5, alpha=0.3, zorder=1)

        for bar, cov in zip(bars, coverages):
            x_pos = bar.get_width() + 1.5
            ax.text(x_pos, bar.get_y() + bar.get_height()/2,
                    f'{cov:.0f}%', va='center', ha='left', fontsize=8,
                    fontweight='bold', color=colors_map['navy'], zorder=5)

        ax.set_xlim(0, 110)
        ax.set_xlabel('Verification Coverage (%)', fontsize=9, color=colors_map['navy'], fontweight='bold')
        ax.set_title('Work Product Verification Coverage', fontsize=11, fontweight='bold',
                     color=colors_map['navy'], pad=12)
        ax.tick_params(axis='y', labelsize=7.5)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_verif_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Verification coverage chart generation failed: {e}")
        return None


# ═══════════════════════════════════════════════════
#  NEW ENHANCED VISUALIZATIONS
# ═══════════════════════════════════════════════════

def generate_radar_chart(layers_data, layer_labels):
    """
    Generate a radar/spider chart showing compliance across all algorithm layers.
    Provides an at-a-glance multi-dimensional compliance overview.
    Returns temp PNG path.
    """
    if not HAS_MATPLOTLIB or not HAS_NUMPY or not layers_data:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()

        # Collect data
        labels_list = []
        values = []
        for key in layer_labels:
            if key in layers_data:
                labels_list.append(layer_labels[key].replace(' ', '\n'))
                values.append(layers_data[key].get("score", 0))

        if len(labels_list) < 3:
            return None

        N = len(labels_list)
        angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
        values_closed = values + [values[0]]
        angles_closed = angles + [angles[0]]

        fig, ax = plt.subplots(figsize=(5.5, 5.5), subplot_kw=dict(polar=True))
        fig.patch.set_facecolor(colors_map['bg'])

        # Draw reference circles
        for pct in [25, 50, 75, 100]:
            circle = np.full(len(angles_closed), pct)
            ax.plot(angles_closed, circle, color=colors_map['light_grey'],
                    linewidth=0.6, linestyle='--', alpha=0.6)
            ax.text(angles[0], pct + 2, f'{pct}%', fontsize=6,
                    color=colors_map['grey'], ha='center')

        # Threshold zone (60-85)
        theta_fill = np.linspace(0, 2 * np.pi, 100)
        ax.fill_between(theta_fill, 60, 85, alpha=0.06, color=colors_map['yellow'])

        # Data area
        ax.fill(angles_closed, values_closed, alpha=0.18, color=colors_map['blue'])
        ax.plot(angles_closed, values_closed, color=colors_map['blue'],
                linewidth=2.2, marker='o', markersize=6, markerfacecolor=colors_map['navy'],
                markeredgecolor='white', markeredgewidth=1.5)

        # Score labels at each point — offset outward to avoid overlap
        for angle, val, label in zip(angles, values, labels_list):
            color = colors_map['green'] if val >= 80 else (colors_map['yellow'] if val >= 60 else colors_map['red'])
            ax.text(angle, val + 10, f'{val:.0f}%', ha='center', va='center',
                    fontsize=8.5, fontweight='bold', color=color)

        ax.set_xticks(angles)
        ax.set_xticklabels(labels_list, fontsize=7.5, color=colors_map['navy'])
        ax.set_ylim(0, 115)
        ax.set_yticks([])
        ax.tick_params(axis='x', pad=12)  # Push axis labels outward
        ax.spines['polar'].set_color(colors_map['light_grey'])
        ax.set_title('Multi-Layer Compliance Radar', fontsize=12, fontweight='bold',
                     color=colors_map['navy'], pad=25, y=1.10)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_radar_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Radar chart generation failed: {e}")
        return None


def generate_risk_heatmap(gaps, layers_data):
    """
    Generate a risk heatmap — Gap categories vs. severity levels.
    Uses seaborn if available, falls back to matplotlib.
    Returns temp PNG path.
    """
    if not HAS_MATPLOTLIB or not HAS_NUMPY:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()

        # Build matrix from gap data
        categories = []
        sev_levels = ['CRITICAL', 'MAJOR', 'WARNING', 'INFO']
        matrix_data = []

        gap_items = gaps if isinstance(gaps, list) else gaps.get("gaps", [])
        if not gap_items:
            return None

        for g in gap_items[:10]:  # Top 10 gap categories
            key = g.get("key", "Unknown")[:25]
            categories.append(key)
            row = []
            findings = g.get("findings", [])
            sev_dist = {}
            for f in findings:
                s = f.get("severity", "INFO").upper()
                sev_dist[s] = sev_dist.get(s, 0) + 1
            for s in sev_levels:
                row.append(sev_dist.get(s, 0))
            matrix_data.append(row)

        if not matrix_data:
            return None

        matrix = np.array(matrix_data, dtype=float)

        fig, ax = plt.subplots(figsize=(5.5, max(2.5, len(categories) * 0.4 + 0.8)))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])

        if HAS_SEABORN:
            cmap = sns.color_palette("YlOrRd", as_cmap=True)
        else:
            cmap = plt.cm.YlOrRd

        im = ax.imshow(matrix, cmap=cmap, aspect='auto', vmin=0)

        # Annotations
        for i in range(len(categories)):
            for j in range(len(sev_levels)):
                val = int(matrix[i, j])
                if val > 0:
                    text_color = 'white' if val > max(1, matrix.max() * 0.5) else colors_map['dark_text']
                    ax.text(j, i, str(val), ha='center', va='center',
                            fontsize=9, fontweight='bold', color=text_color)

        ax.set_xticks(range(len(sev_levels)))
        ax.set_xticklabels(sev_levels, fontsize=8, color=colors_map['navy'], fontweight='bold')
        ax.set_yticks(range(len(categories)))
        ax.set_yticklabels(categories, fontsize=7.5, color=colors_map['navy'])

        cbar = fig.colorbar(im, ax=ax, shrink=0.8, aspect=30, pad=0.02)
        cbar.set_label('Finding Count', fontsize=8, color=colors_map['grey'])
        cbar.ax.tick_params(labelsize=7)

        ax.set_title('Risk Heatmap — Gap Category × Severity', fontsize=11,
                     fontweight='bold', color=colors_map['navy'], pad=12)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_heatmap_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Risk heatmap generation failed: {e}")
        return None


def generate_waterfall_chart(layers_data, layer_labels, layer_weights, overall_score):
    """
    Generate a waterfall chart showing how each layer contributes to the
    final weighted compliance score. Each bar shows the layer's weighted
    contribution, building up to the total.
    Returns temp PNG path.
    """
    if not HAS_MATPLOTLIB or not HAS_NUMPY or not layers_data:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()

        # Calculate contributions
        items = []
        for key in layer_labels:
            if key in layers_data:
                score = layers_data[key].get("score", 0)
                weight = layer_weights.get(key, 0)
                contribution = score * weight
                items.append((layer_labels[key], contribution, score, weight))

        if not items:
            return None

        labels = [it[0] for it in items] + ['Total']
        contributions = [it[1] for it in items]
        total = sum(contributions)

        fig, ax = plt.subplots(figsize=(7.0, 3.5))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])

        # Running total for waterfall positions
        running = 0
        bar_bottoms = []
        bar_heights = []
        bar_colors_list = []

        for contrib in contributions:
            bar_bottoms.append(running)
            bar_heights.append(contrib)
            if contrib >= 12:
                bar_colors_list.append(colors_map['green'])
            elif contrib >= 6:
                bar_colors_list.append(colors_map['blue'])
            else:
                bar_colors_list.append(colors_map['yellow'])
            running += contrib

        # Total bar (from 0)
        bar_bottoms.append(0)
        bar_heights.append(total)
        bar_colors_list.append(colors_map['navy'])

        x = np.arange(len(labels))
        bars = ax.bar(x, bar_heights, bottom=bar_bottoms, color=bar_colors_list,
                      width=0.55, edgecolor='white', linewidth=1.0, alpha=0.9, zorder=3)

        # Connector lines between waterfall bars
        for i in range(len(contributions)):
            connector_y = bar_bottoms[i] + bar_heights[i]
            if i < len(contributions) - 1:
                ax.plot([x[i] + 0.275, x[i+1] - 0.275], [connector_y, connector_y],
                        color=colors_map['grey'], linewidth=0.8, linestyle=':', zorder=2)

        # Value labels — position above bar if bar is too short for text inside
        for i, bar in enumerate(bars):
            val = bar_heights[i]
            bar_top = bar_bottoms[i] + val
            if i < len(items):
                sc, wt = items[i][2], items[i][3]
                label_text = f'{val:.1f}  ({sc:.0f}% × {wt:.0%})'
            else:
                label_text = f'{val:.1f}%'
            # Place label above bar if bar segment is too narrow
            if val < (total * 0.08) and i < len(items):
                ax.text(bar.get_x() + bar.get_width()/2, bar_top + total * 0.02, label_text,
                        ha='center', va='bottom', fontsize=7.5, fontweight='bold',
                        color=colors_map['dark_text'], zorder=5)
            else:
                y_pos = bar_bottoms[i] + val / 2
                ax.text(bar.get_x() + bar.get_width()/2, y_pos, label_text,
                        ha='center', va='center', fontsize=7.5, fontweight='bold',
                        color='white' if i == len(items) else colors_map['dark_text'], zorder=5)

        ax.set_xticks(x)
        ax.set_xticklabels(labels, fontsize=7, color=colors_map['navy'], rotation=30, ha='right')
        ax.set_ylabel('Weighted Score Contribution', fontsize=9, color=colors_map['navy'], fontweight='bold')
        ax.set_title('Compliance Score Waterfall — Layer Contributions', fontsize=11,
                     fontweight='bold', color=colors_map['navy'], pad=12)

        # Reference line at overall score
        ax.axhline(y=total, color=colors_map['navy'], linewidth=1.0, linestyle='--', alpha=0.5, zorder=1)

        ax.tick_params(axis='y', labelsize=7.5, colors=colors_map['grey'])
        ax.spines['left'].set_color(colors_map['light_grey'])
        ax.spines['bottom'].set_color(colors_map['light_grey'])

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_waterfall_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Waterfall chart generation failed: {e}")
        return None


def generate_traceability_slope_chart(trace_mx, layers_data):
    """
    Generate a slope/alluvial chart showing requirement-to-evidence traceability.
    Left column: Standard requirements. Right column: Artifact evidence.
    Lines connect traced pairs (green=linked, red=orphan).
    Returns temp PNG path.
    """
    if not HAS_MATPLOTLIB or not HAS_NUMPY:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()

        trace_cov = trace_mx.get("coverage_summary", {})
        total_reqs = trace_cov.get("total_requirements", 0)
        traced = trace_cov.get("traced_requirements", 0)
        orphan_reqs = trace_cov.get("orphan_requirements_count", 0)
        orphan_arts = trace_cov.get("orphan_artifacts_count", 0)

        if total_reqs == 0 and traced == 0:
            return None

        fig, ax = plt.subplots(figsize=(6.5, 4.5))
        fig.patch.set_facecolor(colors_map['bg'])
        ax.set_facecolor(colors_map['bg'])

        # Left column: Requirements (traced + orphan reqs)
        # Right column: Artifacts (traced + orphan arts)
        left_traced = traced
        left_orphan = orphan_reqs
        right_traced = traced
        right_orphan = orphan_arts

        left_total = left_traced + left_orphan
        right_total = right_traced + right_orphan
        max_total = max(left_total, right_total, 1)

        # Normalize heights
        def y_positions(traced_n, orphan_n, total_max):
            positions = []
            h_traced = (traced_n / total_max) if total_max > 0 else 0
            h_orphan = (orphan_n / total_max) if total_max > 0 else 0
            return h_traced, h_orphan

        lt, lo = y_positions(left_traced, left_orphan, max_total)
        rt, ro = y_positions(right_traced, right_orphan, max_total)

        x_left = 0.15
        x_right = 0.85
        bar_width = 0.12

        # Draw stacked bars — LEFT
        ax.barh(0.5, bar_width, left=x_left - bar_width/2, height=lt,
                color=colors_map['green'], alpha=0.85, zorder=3)
        ax.barh(0.5 + lt, bar_width, left=x_left - bar_width/2, height=lo,
                color=colors_map['red'], alpha=0.85, zorder=3)

        # Draw stacked bars — RIGHT
        ax.barh(0.5, bar_width, left=x_right - bar_width/2, height=rt,
                color=colors_map['green'], alpha=0.85, zorder=3)
        ax.barh(0.5 + rt, bar_width, left=x_right - bar_width/2, height=ro,
                color=colors_map['amber'], alpha=0.85, zorder=3)

        # Flow bands (bezier curves connecting traced portions)
        from matplotlib.patches import FancyBboxPatch
        verts_traced = [
            (x_left + bar_width/2, 0.5),
            (0.5, 0.5),
            (0.5, 0.5),
            (x_right - bar_width/2, 0.5),
            (x_right - bar_width/2, 0.5 + rt),
            (0.5, 0.5 + min(lt, rt)),
            (0.5, 0.5 + min(lt, rt)),
            (x_left + bar_width/2, 0.5 + lt),
        ]
        codes_traced = [Path.MOVETO, Path.CURVE4, Path.CURVE4, Path.LINETO,
                        Path.LINETO, Path.CURVE4, Path.CURVE4, Path.LINETO]
        path_traced = Path(verts_traced, codes_traced)
        patch_traced = mpatches.PathPatch(path_traced, facecolor=colors_map['green'],
                                          alpha=0.15, edgecolor=colors_map['green'],
                                          linewidth=0.8, zorder=2)
        ax.add_patch(patch_traced)

        # Labels
        ax.text(x_left, 0.42, 'STANDARD\nREQUIREMENTS', ha='center', va='top',
                fontsize=9, fontweight='bold', color=colors_map['navy'])
        ax.text(x_right, 0.42, 'ARTIFACT\nEVIDENCE', ha='center', va='top',
                fontsize=9, fontweight='bold', color=colors_map['navy'])

        # Count labels
        ax.text(x_left, 0.5 + lt/2, f'{left_traced}\nTraced', ha='center', va='center',
                fontsize=8, fontweight='bold', color='white', zorder=5)
        if lo > 0.03:
            ax.text(x_left, 0.5 + lt + lo/2, f'{left_orphan}\nOrphan', ha='center', va='center',
                    fontsize=7, fontweight='bold', color='white', zorder=5)

        ax.text(x_right, 0.5 + rt/2, f'{right_traced}\nLinked', ha='center', va='center',
                fontsize=8, fontweight='bold', color='white', zorder=5)
        if ro > 0.03:
            ax.text(x_right, 0.5 + rt + ro/2, f'{right_orphan}\nOrphan', ha='center', va='center',
                    fontsize=7, fontweight='bold', color='white', zorder=5)

        # Coverage percentage
        cov_pct = trace_cov.get("trace_coverage_pct", 0)
        cov_color = colors_map['green'] if cov_pct >= 80 else (colors_map['yellow'] if cov_pct >= 60 else colors_map['red'])
        ax.text(0.5, 0.5 + max(lt, rt)/2, f'{cov_pct}%\ncoverage',
                ha='center', va='center', fontsize=14, fontweight='bold',
                color=cov_color, zorder=5)

        ax.set_xlim(0, 1)
        ax.set_ylim(0.3, 0.5 + max(lt + lo, rt + ro) + 0.15)
        ax.axis('off')

        ax.set_title('Traceability Flow — Requirements ⟷ Evidence', fontsize=12,
                     fontweight='bold', color=colors_map['navy'], pad=15)

        # Legend
        legend_elements = [
            mpatches.Patch(facecolor=colors_map['green'], alpha=0.85, label='Traced / Linked'),
            mpatches.Patch(facecolor=colors_map['red'], alpha=0.85, label='Orphan Requirements'),
            mpatches.Patch(facecolor=colors_map['amber'], alpha=0.85, label='Orphan Artifacts'),
        ]
        ax.legend(handles=legend_elements, loc='lower center', ncol=3, fontsize=7.5,
                  frameon=False, bbox_to_anchor=(0.5, -0.08))

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_slope_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Traceability slope chart generation failed: {e}")
        return None


def generate_safety_decomposition_tree(report):
    """
    Generate a clean V-model safety decomposition tree using Graphviz.
    Shows the ISO 26262 V-model: HARA → Safety Goals → FSR → TSR → HSR/SSR.
    Compact and readable design. Returns temp PNG path.
    """
    if not HAS_GRAPHVIZ:
        return None
    try:
        schema = report.get("discovered_schema", {})

        dot = graphviz.Digraph(
            format='png',
            graph_attr={
                'rankdir': 'TB',
                'bgcolor': CHART_BG,
                'fontname': 'DejaVu Sans',
                'fontsize': '11',
                'nodesep': '0.6',
                'ranksep': '0.55',
                'dpi': '180',
                'pad': '0.4',
                'margin': '0.3',
                'splines': 'ortho',
            },
            node_attr={
                'fontname': 'DejaVu Sans',
                'fontsize': '10',
                'style': 'filled,rounded',
                'shape': 'box',
                'penwidth': '1.8',
                'margin': '0.2,0.1',
            },
            edge_attr={
                'fontname': 'DejaVu Sans',
                'fontsize': '8',
                'color': '#4B5563',
                'penwidth': '1.5',
                'arrowsize': '0.8',
            }
        )

        # Standard root
        std_name = schema.get("standard_name", "ISO 26262")
        dot.node('STD', f'{std_name[:35]}',
                 fillcolor='#0D2B4E', fontcolor='white',
                 shape='box', fontsize='12', penwidth='2.5')

        # V-model hierarchy nodes
        nodes = [
            ('HARA', 'Hazard Analysis\n& Risk Assessment', '#B91C1C', 'white'),
            ('SG',   'Safety Goals',                        '#D46B08', 'white'),
            ('FSC',  'Functional Safety\nConcept',          '#CA8A04', 'white'),
            ('TSC',  'Technical Safety\nConcept',           '#2E6EB5', 'white'),
        ]

        # Left branch (HW) and Right branch (SW)
        hw_nodes = [
            ('HW_REQ', 'HW Safety\nRequirements',   '#166B34', 'white'),
            ('HW_DES', 'HW Design &\nIntegration',  '#166B34', 'white'),
        ]
        sw_nodes = [
            ('SW_REQ', 'SW Safety\nRequirements',   '#6B21A8', 'white'),
            ('SW_DES', 'SW Design &\nUnit Testing', '#6B21A8', 'white'),
        ]

        # Add hierarchy
        dot.edge('STD', 'HARA')
        prev = 'HARA'
        for nid, label, fill, fc in nodes:
            dot.node(nid, label, fillcolor=fill, fontcolor=fc)
            dot.edge(prev, nid)
            prev = nid

        # Branches from TSC
        for nid, label, fill, fc in hw_nodes:
            dot.node(nid, label, fillcolor=fill, fontcolor=fc)
        for nid, label, fill, fc in sw_nodes:
            dot.node(nid, label, fillcolor=fill, fontcolor=fc)

        dot.edge('TSC', 'HW_REQ')
        dot.edge('TSC', 'SW_REQ')
        dot.edge('HW_REQ', 'HW_DES')
        dot.edge('SW_REQ', 'SW_DES')

        # Integration & validation (bottom of V)
        dot.node('INT', 'System Integration\n& Verification', fillcolor='#4B5563', fontcolor='white')
        dot.edge('HW_DES', 'INT')
        dot.edge('SW_DES', 'INT')

        # Safety validation
        dot.node('VAL', 'Safety Validation\n& Assessment', fillcolor='#0D2B4E', fontcolor='white',
                 penwidth='2.5')
        dot.edge('INT', 'VAL')

        # Same-rank constraints for parallel branches
        with dot.subgraph() as s:
            s.attr(rank='same')
            s.node('HW_REQ')
            s.node('SW_REQ')
        with dot.subgraph() as s:
            s.attr(rank='same')
            s.node('HW_DES')
            s.node('SW_DES')

        # Render to temp file
        tmp_base = tempfile.mktemp(prefix='chart_tree_')
        dot.render(tmp_base, cleanup=True)
        result_path = tmp_base + '.png'
        if os.path.isfile(result_path):
            return result_path
        return None
    except Exception as e:
        print(f"    [WARN] Safety decomposition tree generation failed: {e}")
        return None


def generate_compliance_timeline_chart(layers_data, layer_labels, layer_weights):
    """
    Generate a horizontal grouped bar chart comparing raw scores vs. weighted contributions.
    Provides transparency into how the engine weighs different aspects.
    Returns temp PNG path.
    """
    if not HAS_MATPLOTLIB or not HAS_NUMPY or not layers_data:
        return None
    try:
        _apply_professional_style()
        colors_map = _chart_colors()

        labels = []
        raw_scores = []
        weighted_scores = []
        for key in layer_labels:
            if key in layers_data:
                labels.append(layer_labels[key])
                sc = layers_data[key].get("score", 0)
                wt = layer_weights.get(key, 0)
                raw_scores.append(sc)
                weighted_scores.append(sc * wt)

        if not labels:
            return None

        x = np.arange(len(labels))
        width = 0.35

        fig, ax1 = plt.subplots(figsize=(7.0, 3.5))
        fig.patch.set_facecolor(colors_map['bg'])
        ax1.set_facecolor(colors_map['bg'])

        # Raw scores (left axis)
        bars1 = ax1.bar(x - width/2, raw_scores, width, label='Raw Score (%)',
                        color=colors_map['blue'], alpha=0.8, edgecolor='white',
                        linewidth=0.8, zorder=3)

        # Weighted contributions (right axis)
        ax2 = ax1.twinx()
        bars2 = ax2.bar(x + width/2, weighted_scores, width, label='Weighted Contribution',
                        color=colors_map['amber'], alpha=0.8, edgecolor='white',
                        linewidth=0.8, zorder=3)

        # Labels on bars
        for bar, val in zip(bars1, raw_scores):
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                     f'{val:.0f}%', ha='center', va='bottom', fontsize=7,
                     fontweight='bold', color=colors_map['blue'])
        for bar, val in zip(bars2, weighted_scores):
            ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                     f'{val:.1f}', ha='center', va='bottom', fontsize=7,
                     fontweight='bold', color=colors_map['amber'])

        ax1.set_xticks(x)
        ax1.set_xticklabels(labels, fontsize=7, color=colors_map['navy'], rotation=25, ha='right')
        ax1.set_ylabel('Raw Score (%)', fontsize=9, color=colors_map['blue'], fontweight='bold')
        ax2.set_ylabel('Weighted Contribution', fontsize=9, color=colors_map['amber'], fontweight='bold')
        ax1.set_ylim(0, 110)
        ax2.set_ylim(0, max(weighted_scores) * 1.3 if weighted_scores else 30)

        ax1.tick_params(axis='y', labelsize=7.5, colors=colors_map['blue'])
        ax2.tick_params(axis='y', labelsize=7.5, colors=colors_map['amber'])

        # Combined legend
        lines1, labels1 = ax1.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper right',
                   fontsize=7.5, framealpha=0.9)

        ax1.set_title('Score Breakdown — Raw vs. Weighted', fontsize=11,
                      fontweight='bold', color=colors_map['navy'], pad=12)
        ax1.spines['top'].set_visible(False)
        ax2.spines['top'].set_visible(False)

        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='chart_compare_')
        fig.savefig(tmp.name, dpi=CHART_DPI, bbox_inches='tight', facecolor=colors_map['bg'],
                    edgecolor='none')
        plt.close(fig)
        return tmp.name
    except Exception as e:
        print(f"    [WARN] Compliance timeline chart generation failed: {e}")
        return None


def resolve_fields(report):
    project = report.get("project", {})
    if not project:
        art_src = report.get("artifact_source", "")
        project = {"name": os.path.basename(art_src).replace(".json", "").replace("_", " ").title()
                    if art_src else "N/A"}

    std_name = report.get("standard_name",
               report.get("standard",
               report.get("discovered_schema", {}).get("standard_name", "")))
    if not std_name:
        src = report.get("standard_source", "")
        if "26262" in src: std_name = "ISO 26262:2018 — Functional Safety for Road Vehicles"
        elif "21434" in src: std_name = "ISO/SAE 21434 — Cybersecurity Engineering"
        elif "sotif" in src.lower(): std_name = "ISO 21448 — SOTIF"
        elif "aspice" in src.lower(): std_name = "Automotive SPICE"
        else: std_name = os.path.basename(src) if src else "Standard"

    return project, std_name


# ═══════════════════════════════════════════════════
#  TABLE BUILDERS
# ═══════════════════════════════════════════════════

def make_info_row(doc, label, value, table=None):
    """Add a key-value row to an info table, or create one."""
    if table is None:
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.add_row()
    c0, c1 = row.cells[0], row.cells[1]
    c0.width = Inches(2.2)
    c1.width = Inches(4.3)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p0, label, bold=True, size=9, color=NAVY)
    p1 = c1.paragraphs[0]
    add_run(p1, str(value), size=9, color=BLACK)
    set_cell_border(c0, bottom="4 single E5E7EB")
    set_cell_border(c1, bottom="4 single E5E7EB")
    clear_borders(c0)
    set_cell_border(c0, bottom="4 single E5E7EB")
    clear_borders(c1)
    set_cell_border(c1, bottom="4 single E5E7EB")
    return table

def add_bar_chart(doc, data, width=6.5):
    """Visual bar chart. data = [(label, score, weight), ...]"""
    if not data:
        return
    table = doc.add_table(rows=len(data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, score, weight) in enumerate(data):
        row = table.rows[i]
        c0, c1, c2, c3 = row.cells
        c0.width = Inches(1.8)
        c1.width = Inches(3.0)
        c2.width = Inches(0.8)
        c3.width = Inches(0.7)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p0, label, size=8, color=GREY_TEXT)
        p1 = c1.paragraphs[0]
        bar_len = int(score / 100 * 28)
        sc = score_color(score)
        if bar_len > 0:
            add_run(p1, "\u2588" * bar_len, size=8, color=sc)
        if bar_len < 28:
            add_run(p1, "\u2591" * (28 - bar_len), size=8, color=RGBColor(0xE5, 0xE7, 0xEB))
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, f"{score:.0f}%", bold=True, size=8, color=sc)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p3, f"({weight:.0%})", size=7, color=GREY_LIGHT)
    for row in table.rows:
        for cell in row.cells:
            clear_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table

def add_severity_boxes(doc, counts):
    items = [
        ("CRITICAL", counts.get("CRITICAL", 0), "B91C1C"),
        ("MAJOR", counts.get("MAJOR", 0), "D46B08"),
        ("WARNING", counts.get("WARNING", 0), "CA8A04"),
        ("INFO", counts.get("INFO", 0), "6B7280"),
    ]
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, count, bg) in enumerate(items):
        c_top = table.rows[0].cells[i]
        set_cell_shading(c_top, bg)
        p = c_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(count), bold=True, size=18, color=WHITE)
        c_top.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_bot = table.rows[1].cells[i]
        set_cell_shading(c_bot, bg)
        p2 = c_bot.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, label, size=7, color=WHITE)
        c_bot.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top="0 none FFFFFF", bottom="0 none FFFFFF",
                            left="4 single FFFFFF", right="4 single FFFFFF")
    return table

def add_metric_boxes(doc, metrics):
    """metrics = [(label, value), ...]"""
    n = len(metrics)
    if n == 0:
        return
    table = doc.add_table(rows=2, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(metrics):
        cv = table.rows[0].cells[i]
        p = cv.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(value), bold=True, size=13, color=NAVY)
        cv.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cl = table.rows[1].cells[i]
        p2 = cl.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, label, size=7.5, color=GREY_TEXT)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top="4 single E5E7EB", bottom="4 single E5E7EB",
                            left="0 none FFFFFF", right="0 none FFFFFF")


# ═══════════════════════════════════════════════════
#  MAIN REPORT BUILDER
# ═══════════════════════════════════════════════════

def generate_compliance_docx(report: dict, output_path: str):
    """
    Generate a professional, audit-grade compliance assessment report.

    Args:
        report: dict from agnostic_engine.run_compliance_check()
        output_path: where to save the DOCX
    """
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.7)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

    # ── Resolve Fields ──
    project, std_name = resolve_fields(report)
    score = report.get("compliance_score", 0)
    grd = report.get("grade", "N/A")
    gap = report.get("gap_analysis", {})
    layers = report.get("layer_results", {})
    summary = report.get("summary", {})
    schema = report.get("discovered_schema", {})
    risk_level = report.get("target_risk_level", "")
    check_date = report.get("check_date", "")[:10]
    decision = report.get("assessment_decision", {})
    wp_register = report.get("work_product_register", [])
    trace_mx = report.get("traceability_matrix", {})
    wp_ref = report.get("iso26262_wp_reference", {})
    stds_assessed = report.get("standards_assessed", [])
    arts_assessed = report.get("artifacts_assessed", [])

    # Parse grade
    grade_text = grd
    grade_desc = ""
    for sep in ["\u2014", "—", " - ", "-"]:
        if sep in grd:
            grade_text = grd.split(sep)[0].strip()
            grade_desc = grd.split(sep, 1)[1].strip()
            break

    # Document ID
    doc_id = f"CIQ-FSA-{check_date.replace('-', '')}" if check_date else "CIQ-FSA-DRAFT"

    # ── Track temp chart files for cleanup ──
    _chart_files = []

    # ═══════════════════════════════════
    #  COVER PAGE
    # ═══════════════════════════════════

    # Logo
    logo_path = find_logo()
    if logo_path:
        try:
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = lp.add_run()
            run.add_picture(logo_path, width=Inches(2.0))
            doc.add_paragraph("")
        except Exception:
            for _ in range(2):
                doc.add_paragraph("")
    else:
        for _ in range(2):
            doc.add_paragraph("")

    # Brand name
    bp = doc.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(bp, BRAND_TAGLINE, bold=True, size=11, color=DARK_BLUE, italic=True)
    bp.paragraph_format.space_after = Pt(8)

    # Top accent line
    add_hr(doc, color_hex="0D2B4E", weight="12")

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    add_run(title, "FUNCTIONAL SAFETY\nCOMPLIANCE ASSESSMENT REPORT", bold=True, size=26, color=NAVY)
    title.paragraph_format.space_after = Pt(2)

    add_hr(doc, color_hex="0D2B4E", weight="12")
    doc.add_paragraph("")

    # Standard subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, std_name, size=13, color=DARK_BLUE, italic=True)
    sub.paragraph_format.space_after = Pt(20)

    # Document info table
    info_table = None
    info_items = [
        ("Document ID", doc_id),
        ("Revision", "1.0"),
        ("Assessment Date", check_date),
        ("Project", project.get("name", "N/A")),
        ("Standard", std_name),
        ("Target ASIL / Risk Level", risk_level if risk_level else "All levels"),
        ("Overall Score", f"{score}%"),
        ("Assessment Grade", grd),
        ("Verdict", decision.get("verdict", "PENDING")),
    ]
    for label, value in info_items:
        info_table = make_info_row(doc, label, value, info_table)

    doc.add_paragraph("")

    # Verdict banner
    verdict = decision.get("verdict", "PENDING")
    v_table = doc.add_table(rows=1, cols=1)
    v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    v_cell = v_table.rows[0].cells[0]
    v_color = verdict_color(verdict)
    v_hex = str(v_color)
    set_cell_shading(v_cell, v_hex)
    vp = v_cell.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(vp, f"  {verdict}  ", bold=True, size=16, color=WHITE)
    v_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Confidential footer
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(conf, "CONFIDENTIAL — FOR AUTHORIZED RECIPIENTS ONLY", size=9, color=GREY_LIGHT)

    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(gen, f"Generated by Lion of Functional Safety Engine\u2122 v2.1 — {check_date}", size=8, color=GREY_LIGHT)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  DOCUMENT CONTROL
    # ═══════════════════════════════════
    add_heading(doc, "Document Control", level=1)
    add_hr(doc)

    # Revision history
    p = doc.add_paragraph()
    add_run(p, "Revision History", bold=True, size=11, color=DARK_BLUE)

    rev_table = doc.add_table(rows=2, cols=5)
    rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Rev.", "Date", "Author", "Reviewer", "Description"]):
        cell = rev_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=8, color=WHITE)
    # Data row
    rev_data = ["1.0", check_date, "ComplianceIQ Engine", "—", "Initial automated assessment"]
    for i, val in enumerate(rev_data):
        p = rev_table.rows[1].cells[i].paragraphs[0]
        add_run(p, val, size=8, color=BLACK)

    doc.add_paragraph("")

    # Distribution
    p = doc.add_paragraph()
    add_run(p, "Distribution", bold=True, size=11, color=DARK_BLUE)
    p = doc.add_paragraph()
    add_run(p, "This document is classified as CONFIDENTIAL and is intended for the project team, "
            "safety managers, and designated assessors only. Redistribution requires written authorization.",
            size=9, color=GREY_TEXT)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_run(p, "Abbreviations", bold=True, size=11, color=DARK_BLUE)

    abbrev_table = doc.add_table(rows=1, cols=2)
    abbrev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Abbreviation", "Definition"]):
        cell = abbrev_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG2)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=8, color=WHITE)

    abbreviations = [
        ("ASIL", "Automotive Safety Integrity Level"),
        ("FSA", "Functional Safety Assessment"),
        ("FSC", "Functional Safety Concept"),
        ("FMEA", "Failure Mode and Effects Analysis"),
        ("FTA", "Fault Tree Analysis"),
        ("HARA", "Hazard Analysis and Risk Assessment"),
        ("HSI", "Hardware-Software Interface"),
        ("LFM", "Latent Fault Metric"),
        ("PMHF", "Probabilistic Metric for Random Hardware Failures"),
        ("SPFM", "Single-Point Fault Metric"),
        ("TSC", "Technical Safety Concept"),
        ("TSR", "Technical Safety Requirements"),
        ("V&V", "Verification and Validation"),
    ]
    for i, (abbr, defn) in enumerate(abbreviations):
        row = abbrev_table.add_row()
        p0 = row.cells[0].paragraphs[0]
        add_run(p0, abbr, bold=True, size=8, color=BLACK)
        p1 = row.cells[1].paragraphs[0]
        add_run(p1, defn, size=8, color=GREY_TEXT)
        style_data_row(row, i)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  TABLE OF CONTENTS
    # ═══════════════════════════════════
    add_heading(doc, "Table of Contents", level=1)
    add_hr(doc)

    toc_items = [
        ("1", "Executive Summary"),
        ("2", "Assessment Scope & Methodology"),
        ("3", "Work Product Register"),
        ("4", "Compliance Assessment Dashboard"),
        ("5", "Requirements Compliance Matrix"),
        ("6", "Traceability Matrix"),
        ("7", "Gap Analysis"),
        ("8", "Method & Practice Audit"),
        ("9", "Risk Assessment"),
        ("10", "Corrective Action Register"),
        ("11", "Confirmation Review"),
        ("12", "Verification Review"),
        ("13", "Assessment Decision"),
        ("A", "Appendix A: Detailed Findings"),
        ("B", "Appendix B: Algorithm Layer Details"),
        ("C", "Appendix C: Glossary & References"),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        add_run(p, f"{num}.  " if num.isdigit() else f"{num}.  ", bold=True, size=10, color=NAVY)
        add_run(p, title_text, size=10, color=BLACK)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════
    add_heading(doc, "1.  Executive Summary", level=1)
    add_hr(doc)

    # Verdict summary
    p = doc.add_paragraph()
    add_run(p, "Assessment Verdict: ", bold=True, size=11, color=NAVY)
    add_run(p, verdict, bold=True, size=11, color=verdict_color(verdict))
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    add_run(p, decision.get("description", ""), size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(6)

    if decision.get("conditions"):
        p = doc.add_paragraph()
        add_run(p, "Conditions for Acceptance:", bold=True, size=9.5, color=ACCENT)
        for cond in decision["conditions"]:
            p = doc.add_paragraph()
            add_run(p, f"    \u2022  {cond}", size=9, color=BLACK)
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")

    # Key metrics row
    add_metric_boxes(doc, [
        ("Compliance Score", f"{score}%"),
        ("Grade", grade_text),
        ("Total Findings", str(summary.get("total_findings", 0))),
        ("Critical", str(gap.get("severity_counts", {}).get("CRITICAL", 0))),
        ("Risk Score", str(summary.get("risk_score", 0))),
    ])

    doc.add_paragraph("")

    # Top recommendations
    p = doc.add_paragraph()
    add_run(p, "Priority Actions", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    recs = summary.get("top_recommendations", gap.get("recommendations", []))
    for i, rec in enumerate(recs[:5], 1):
        if "[HIGH]" in rec:
            pri_color = RED_CRIT
        elif "[MEDIUM]" in rec:
            pri_color = ACCENT
        else:
            pri_color = GREY_TEXT
        display = rec.replace("[HIGH]", "").replace("[MEDIUM]", "").replace("[LOW]", "").strip()
        p = doc.add_paragraph()
        add_run(p, f"  {i}. ", bold=True, size=9, color=BLACK)
        add_run(p, display, size=9, color=BLACK)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  2. ASSESSMENT SCOPE & METHODOLOGY
    # ═══════════════════════════════════
    add_heading(doc, "2.  Assessment Scope & Methodology", level=1)
    add_hr(doc)

    # 2.1 Scope
    p = doc.add_paragraph()
    add_run(p, "2.1  Scope of Assessment", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, "This report presents the results of an automated compliance assessment of the submitted "
            "work products against the applicable requirements of ", size=9.5, color=BLACK)
    add_run(p, std_name, bold=True, size=9.5, color=BLACK)
    add_run(p, ". The assessment covers all applicable parts and clauses within the declared scope.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    # 2.2 Standards
    p = doc.add_paragraph()
    add_run(p, "2.2  Normative Reference", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    # Count parts assessed
    n_parts = len(stds_assessed) if stds_assessed else 0
    parts_note = f" ({n_parts} parts assessed)" if n_parts > 1 else ""

    p = doc.add_paragraph()
    add_run(p, f"  {std_name}{parts_note}", size=9.5, color=BLACK)

    doc.add_paragraph("")

    # 2.3 Work products
    p = doc.add_paragraph()
    add_run(p, "2.3  Work Products Evaluated", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, f"A total of {len(arts_assessed)} work product(s) were submitted for assessment.",
            size=9.5, color=BLACK)

    doc.add_paragraph("")

    # 2.4 Methodology
    p = doc.add_paragraph()
    add_run(p, "2.4  Assessment Methodology", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, "The assessment employs an 8-layer algorithmic analysis pipeline. Each layer examines "
            "a different dimension of compliance, and results are combined using risk-weighted ensemble "
            "scoring to produce the final compliance determination.", size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(6)

    method_table = doc.add_table(rows=1, cols=4)
    method_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Layer", "Technique", "Weight", "Purpose"]):
        cell = method_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=7.5, color=WHITE)

    method_info = [
        ("1. Node Coverage", "Decision Tree", "20%", "Verify each requirement clause has corresponding content"),
        ("2. Content Alignment", "TF-IDF + Cosine", "15%", "Measure textual similarity between requirements and artifacts"),
        ("3. Semantic Matching", "Ratcliff-Obershelp", "10%", "Fuzzy semantic matching for paraphrased content"),
        ("4. Concept Coverage", "Set Analysis", "10%", "Verify all required concepts/topics are addressed"),
        ("5. Reference Integrity", "Graph BFS", "15%", "Validate cross-references and dependency links"),
        ("6. Method Audit", "Risk-Aware Rules", "10%", "Check required methods/techniques per risk level"),
        ("7. Traceability Chain", "Directed Graph", "20%", "Verify bidirectional traceability paths"),
        ("8. Gap Analysis", "Bayesian Risk", "—", "Ensemble risk scoring and gap prioritization"),
    ]
    for idx, (layer, tech, weight, purpose) in enumerate(method_info):
        row = method_table.add_row()
        for j, val in enumerate([layer, tech, weight, purpose]):
            p = row.cells[j].paragraphs[0]
            bld = j == 0
            add_run(p, val, bold=bld, size=7.5, color=BLACK)
        style_data_row(row, idx)

    doc.add_paragraph("")

    # 2.5 Assessment Criteria
    p = doc.add_paragraph()
    add_run(p, "2.5  Assessment Criteria", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    criteria_table = doc.add_table(rows=1, cols=3)
    criteria_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Verdict", "Criteria", "Action Required"]):
        cell = criteria_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=8, color=WHITE)

    criteria_data = [
        ("COMPLIANT", "Score \u2265 85%, 0 critical, \u2264 2 major findings",
         "Minor observations only — no corrective action required"),
        ("CONDITIONALLY\nCOMPLIANT", "Score \u2265 60%, \u2264 2 critical findings",
         "Corrective actions required within agreed timeline"),
        ("NOT COMPLIANT", "Score < 60% or > 2 critical findings",
         "Comprehensive remediation and re-assessment required"),
    ]
    for idx, (v, c, a) in enumerate(criteria_data):
        row = criteria_table.add_row()
        p0 = row.cells[0].paragraphs[0]
        add_run(p0, v, bold=True, size=8, color=verdict_color(v.replace("\n", " ")))
        p1 = row.cells[1].paragraphs[0]
        add_run(p1, c, size=8, color=BLACK)
        p2 = row.cells[2].paragraphs[0]
        add_run(p2, a, size=8, color=GREY_TEXT)
        style_data_row(row, idx)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  3. WORK PRODUCT REGISTER
    # ═══════════════════════════════════
    add_heading(doc, "3.  Work Product Register", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "The following work products were submitted and evaluated. Each is mapped to the "
            "applicable standard part(s) with an individual compliance status.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    if wp_register:
        wp_table = doc.add_table(rows=1, cols=6)
        wp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["#", "Work Product", "Format", "Mapped Parts", "Findings", "Status"]):
            cell = wp_table.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=7.5, color=WHITE)

        for idx, wp in enumerate(wp_register):
            row = wp_table.add_row()
            vals = [
                str(idx + 1),
                (wp.get("name", "") or wp.get("id", ""))[:40],
                wp.get("format", "—"),
                ", ".join(p.replace("part_", "Part ") for p in wp.get("mapped_parts", []))[:30] or "—",
                f"{wp.get('critical_count', 0)}C / {wp.get('major_count', 0)}M",
                wp.get("status", "—"),
            ]
            for j, val in enumerate(vals):
                p = row.cells[j].paragraphs[0]
                if j == 5:  # Status column
                    add_run(p, val, bold=True, size=7.5, color=status_color(val))
                else:
                    add_run(p, val, size=7.5, color=BLACK)
            style_data_row(row, idx)
    else:
        p = doc.add_paragraph()
        add_run(p, "No structured work product register data available.", italic=True,
                size=9, color=GREY_TEXT)

    doc.add_paragraph("")

    # ISO 26262 Work Product Reference
    if wp_ref:
        p = doc.add_paragraph()
        add_run(p, "ISO 26262 Work Product Reference", bold=True, size=11, color=DARK_BLUE)
        p.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        add_run(p, "The following table lists key work products required by ISO 26262 for reference. "
                "Work products marked with evidence in the register above indicate coverage.",
                size=9, color=GREY_TEXT)
        p.paragraph_format.space_after = Pt(4)

        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Standard Part", "Required Work Products"]):
            cell = ref_table.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG2)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=8, color=WHITE)

        for idx, (part, wps) in enumerate(wp_ref.items()):
            row = ref_table.add_row()
            p0 = row.cells[0].paragraphs[0]
            add_run(p0, part, bold=True, size=7.5, color=NAVY)
            p1 = row.cells[1].paragraphs[0]
            add_run(p1, "\n".join(wps), size=7, color=BLACK)
            style_data_row(row, idx)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  4. COMPLIANCE ASSESSMENT DASHBOARD
    # ═══════════════════════════════════
    add_heading(doc, "4.  Compliance Assessment Dashboard", level=1)
    add_hr(doc)

    # Score and Grade display
    dash_table = doc.add_table(rows=1, cols=2)
    dash_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sc = score_color(score)

    # Score cell
    c_score = dash_table.rows[0].cells[0]
    c_score.width = Inches(3.2)
    p = c_score.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    add_run(p, f"{score}%", bold=True, size=30, color=sc)
    p2 = c_score.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    add_run(p2, "Compliance Score", size=9, color=GREY_TEXT)

    # Grade cell
    c_grade = dash_table.rows[0].cells[1]
    c_grade.width = Inches(3.2)
    p = c_grade.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    add_run(p, grade_text, bold=True, size=30, color=sc)
    p2 = c_grade.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    add_run(p2, grade_desc if grade_desc else "Assessment Grade", size=9, color=sc)

    for row in dash_table.rows:
        for cell in row.cells:
            clear_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")  # Extra spacing before severity boxes

    # Severity — use donut chart if matplotlib available, else boxes
    sev_counts = gap.get("severity_counts", {})

    sev_chart_path = generate_severity_pie(sev_counts)
    if sev_chart_path:
        _chart_files.append(sev_chart_path)
        # Side-by-side: severity chart and severity boxes
        sev_p = doc.add_paragraph()
        sev_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sev_p.add_run()
        run.add_picture(sev_chart_path, width=Inches(2.8))
        doc.add_paragraph("")
    add_severity_boxes(doc, sev_counts)

    doc.add_paragraph("")

    # Layer performance — use matplotlib chart if available
    layer_weights = {
        "node_coverage": 0.20, "content_alignment": 0.15,
        "semantic_matching": 0.10, "concept_coverage": 0.10,
        "reference_integrity": 0.15, "method_audit": 0.10,
        "traceability_chain": 0.20,
    }
    layer_labels = {
        "node_coverage": "Node Coverage",
        "content_alignment": "Content Alignment",
        "semantic_matching": "Semantic Matching",
        "concept_coverage": "Concept Coverage",
        "reference_integrity": "Reference Integrity",
        "method_audit": "Method Audit",
        "traceability_chain": "Traceability Chain",
    }

    p = doc.add_paragraph()
    add_run(p, "Algorithm Layer Performance", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(6)

    layer_chart_path = generate_layer_chart(layers, layer_labels, layer_weights)
    if layer_chart_path:
        _chart_files.append(layer_chart_path)
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run()
        run.add_picture(layer_chart_path, width=Inches(5.8))
    else:
        # Fallback: text-based bar chart
        bar_data = []
        for key, label in layer_labels.items():
            if key in layers:
                bar_data.append((label, layers[key].get("score", 0), layer_weights.get(key, 0)))
        add_bar_chart(doc, bar_data)

    doc.add_paragraph("")

    # Key metrics
    nc_layer = layers.get("node_coverage", {})
    cc_layer = layers.get("concept_coverage", {})
    add_metric_boxes(doc, [
        ("Total Findings", str(summary.get("total_findings", 0))),
        ("Risk Score", str(summary.get("risk_score", 0))),
        ("Requirements", f"{nc_layer.get('covered_nodes', 0):.0f}/{nc_layer.get('applicable_nodes', 0)}"),
        ("Concepts", f"{cc_layer.get('covered', 0)}/{cc_layer.get('required', 0)}"),
    ])

    doc.add_paragraph("")

    # ── NEW: Radar Chart ──
    p = doc.add_paragraph()
    add_run(p, "4.1  Multi-Layer Compliance Radar", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, "The radar chart provides an at-a-glance view of compliance across all "
            "algorithm dimensions. Areas inside the amber zone (60–85%) indicate layers "
            "requiring attention.", size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(6)

    radar_chart_path = generate_radar_chart(layers, layer_labels)
    if radar_chart_path:
        _chart_files.append(radar_chart_path)
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rp.add_run()
        run.add_picture(radar_chart_path, width=Inches(4.2))

    doc.add_paragraph("")

    # ── NEW: Waterfall Chart ──
    p = doc.add_paragraph()
    add_run(p, "4.2  Score Composition Waterfall", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, "Each algorithm layer contributes a weighted portion to the overall compliance "
            "score. The waterfall chart below decomposes the total score into individual "
            "layer contributions, showing how each dimension builds toward the final result.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(6)

    waterfall_chart_path = generate_waterfall_chart(layers, layer_labels, layer_weights, score)
    if waterfall_chart_path:
        _chart_files.append(waterfall_chart_path)
        wp = doc.add_paragraph()
        wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = wp.add_run()
        run.add_picture(waterfall_chart_path, width=Inches(5.8))

    doc.add_paragraph("")

    # ── NEW: Raw vs Weighted Comparison ──
    p = doc.add_paragraph()
    add_run(p, "4.3  Raw Score vs. Weighted Contribution", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, "This comparison reveals how weighting affects the final score. A layer with "
            "a high raw score but low weight contributes less than a lower-scoring layer "
            "with higher weight — highlighting the engine's prioritization of structural "
            "compliance factors.", size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(6)

    compare_chart_path = generate_compliance_timeline_chart(layers, layer_labels, layer_weights)
    if compare_chart_path:
        _chart_files.append(compare_chart_path)
        ccp = doc.add_paragraph()
        ccp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ccp.add_run()
        run.add_picture(compare_chart_path, width=Inches(5.8))

    doc.add_page_break()

    # ═══════════════════════════════════
    #  5. REQUIREMENTS COMPLIANCE MATRIX
    # ═══════════════════════════════════
    add_heading(doc, "5.  Requirements Compliance Matrix", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "Each requirement clause from the standard is assessed against the submitted work products. "
            "Status categories are: ", size=9, color=BLACK)
    add_run(p, "Compliant", bold=True, size=9, color=GREEN_OK)
    add_run(p, " (content verified), ", size=9, color=BLACK)
    add_run(p, "Partially Compliant", bold=True, size=9, color=YELLOW_WARN)
    add_run(p, " (exists but incomplete), ", size=9, color=BLACK)
    add_run(p, "Not Compliant", bold=True, size=9, color=RED_CRIT)
    add_run(p, " (insufficient coverage), ", size=9, color=BLACK)
    add_run(p, "Missing", bold=True, size=9, color=RED_CRIT)
    add_run(p, " (no artifact), ", size=9, color=BLACK)
    add_run(p, "N/A", bold=True, size=9, color=GREY_LIGHT)
    add_run(p, " (not applicable).", size=9, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    nc = layers.get("node_coverage", {})
    ca = layers.get("content_alignment", {})
    ca_scores = ca.get("scores", {})

    sorted_findings = sorted(nc.get("findings", []),
                             key=lambda f: (f.get("group", ""), f.get("node_id", "")))

    missing_ids = set()
    partial_ids = set()

    if sorted_findings:
        matrix = doc.add_table(rows=1, cols=5)
        matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
        matrix.autofit = False

        # Set proportional column widths: Part(0.65), Clause(0.75), Title(3.6), Status(1.0), Align(0.6)
        col_widths = [Inches(0.65), Inches(0.75), Inches(3.6), Inches(1.0), Inches(0.6)]
        for i, h in enumerate(["Part", "Clause", "Title", "Status", "Align %"]):
            cell = matrix.rows[0].cells[i]
            cell.width = col_widths[i]
            set_cell_shading(cell, HDR_BG)
            set_cell_margins(cell, top=30, bottom=30, left=50, right=50)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=7, color=WHITE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for idx, f in enumerate(sorted_findings[:100]):
            nid = f.get("node_id", "")
            group = f.get("group", "").replace("part_", "Part ")
            title_text = f.get("title", "")[:80]  # Allow longer titles
            ftype = f.get("type", "")

            if ftype == "MISSING_COVERAGE":
                st_text, st_color = "Missing", RED_CRIT
                missing_ids.add(nid)
                content_pct = "—"
            elif ftype == "LOW_CONTENT_ALIGNMENT":
                st_text, st_color = "Not Compliant", RED_CRIT
                missing_ids.add(nid)
                pct = ca_scores.get(nid, "—")
                content_pct = f"{pct}%" if pct != "—" else "—"
            else:
                st_text, st_color = "Partially Compliant", YELLOW_WARN
                partial_ids.add(nid)
                pct = ca_scores.get(nid, "—")
                content_pct = f"{pct}%" if pct != "—" else "—"

            row = matrix.add_row()
            cells = row.cells
            for ci, cw in enumerate(col_widths):
                cells[ci].width = cw
                set_cell_margins(cells[ci], top=20, bottom=20, left=50, right=50)

            p0 = cells[0].paragraphs[0]; add_run(p0, group, size=6.5, color=GREY_TEXT)
            p1 = cells[1].paragraphs[0]; add_run(p1, nid, bold=True, size=6.5, color=BLACK)
            p2 = cells[2].paragraphs[0]; add_run(p2, title_text, size=6.5, color=BLACK)
            p3 = cells[3].paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p3, st_text, bold=True, size=6.5, color=st_color)
            p4 = cells[4].paragraphs[0]; p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p4, content_pct, size=6.5, color=BLACK)

            style_data_row(row, idx)

        if len(sorted_findings) > 100:
            p = doc.add_paragraph()
            add_run(p, f"Showing 100 of {len(sorted_findings)} findings. See Appendix A for full listing.",
                    italic=True, size=8, color=GREY_TEXT)
    else:
        p = doc.add_paragraph()
        add_run(p, "\u2714  All requirement clauses are fully covered.", bold=True,
                size=10, color=GREEN_OK)

    # Summary
    total_applicable = nc.get("applicable_nodes", 0)
    n_covered = total_applicable - len(missing_ids) - len(partial_ids)
    doc.add_paragraph("")
    sum_table = doc.add_table(rows=1, cols=4)
    sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val, clr) in enumerate([
        ("Compliant", str(n_covered), GREEN_OK),
        ("Partial", str(len(partial_ids)), YELLOW_WARN),
        ("Non-Compliant", str(len(missing_ids)), RED_CRIT),
        ("Total Applicable", str(total_applicable), NAVY),
    ]):
        cell = sum_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f"{val}\n", bold=True, size=12, color=clr)
        add_run(p, label, size=7.5, color=GREY_TEXT)
        set_cell_border(cell, top="4 single E5E7EB", bottom="4 single E5E7EB",
                        left="0 none FFFFFF", right="0 none FFFFFF")

    doc.add_page_break()

    # ═══════════════════════════════════
    #  6. TRACEABILITY MATRIX
    # ═══════════════════════════════════
    add_heading(doc, "6.  Traceability Matrix", level=1)
    add_hr(doc)

    trace_cov = trace_mx.get("coverage_summary", {})

    p = doc.add_paragraph()
    add_run(p, "Bidirectional traceability is a core requirement of ISO 26262. Each safety requirement "
            "must trace forward to implementation evidence and backward to its parent safety goal. "
            "This section evaluates the completeness of these traceability chains.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    # Traceability summary metrics
    add_metric_boxes(doc, [
        ("Total Requirements", str(trace_cov.get("total_requirements", 0))),
        ("Traced", str(trace_cov.get("traced_requirements", 0))),
        ("Coverage", f"{trace_cov.get('trace_coverage_pct', 0)}%"),
        ("Orphan Reqs", str(trace_cov.get("orphan_requirements_count", 0))),
        ("Orphan Artifacts", str(trace_cov.get("orphan_artifacts_count", 0))),
    ])

    doc.add_paragraph("")

    # Traceability chart
    trace_chart_path = generate_trace_coverage_chart(trace_cov)
    if trace_chart_path:
        _chart_files.append(trace_chart_path)
        tcp = doc.add_paragraph()
        tcp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tcp.add_run()
        run.add_picture(trace_chart_path, width=Inches(4.0))
        doc.add_paragraph("")

    # ── NEW: Traceability Slope/Flow Chart ──
    slope_chart_path = generate_traceability_slope_chart(trace_mx, layers)
    if slope_chart_path:
        _chart_files.append(slope_chart_path)
        p = doc.add_paragraph()
        add_run(p, "Requirement–Evidence Flow", bold=True, size=10, color=DARK_BLUE)
        p.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        add_run(p, "The flow diagram below visualizes the traceability coverage between "
                "standard requirements and artifact evidence. Green flows indicate "
                "successfully traced links; red/amber indicate orphan items requiring action.",
                size=9, color=BLACK)
        p.paragraph_format.space_after = Pt(6)
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sp.add_run()
        run.add_picture(slope_chart_path, width=Inches(5.2))
        doc.add_paragraph("")

    # ── NEW: Safety Goal Decomposition Tree ──
    tree_chart_path = generate_safety_decomposition_tree(report)
    if tree_chart_path:
        _chart_files.append(tree_chart_path)
        p = doc.add_paragraph()
        add_run(p, "Safety Requirements Decomposition", bold=True, size=10, color=DARK_BLUE)
        p.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        add_run(p, "The decomposition tree shows the hierarchical relationship between "
                "the standard's structure and its constituent work products. This view "
                "supports auditor understanding of the safety argument structure.",
                size=9, color=BLACK)
        p.paragraph_format.space_after = Pt(6)
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tp.add_run()
        try:
            run.add_picture(tree_chart_path, width=Inches(5.5))
        except Exception:
            pass
        doc.add_paragraph("")

    # Traceability chain layer results
    tc = layers.get("traceability_chain", {})
    tc_score = tc.get("score", 0)
    tc_req = tc.get("required", 0)
    tc_sat = tc.get("satisfied", 0)

    p = doc.add_paragraph()
    add_run(p, "6.1  Traceability Chain Analysis", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, f"The standard defines ", size=9.5)
    add_run(p, f"{tc_req}", bold=True, size=9.5)
    add_run(p, f" required traceability paths. ", size=9.5)
    add_run(p, f"{tc_sat}", bold=True, size=9.5)
    add_run(p, f" are satisfied (", size=9.5)
    add_run(p, f"{tc_score:.0f}%", bold=True, size=9.5, color=score_color(tc_score))
    add_run(p, f").", size=9.5)

    # Missing traces
    tc_findings = tc.get("findings", [])
    trace_findings = [f for f in tc_findings if f.get("type") != "NO_TRACEABILITY_MODEL"]
    if trace_findings:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        add_run(p, "6.2  Missing Traceability Links", bold=True, size=11, color=DARK_BLUE)
        p.paragraph_format.space_after = Pt(4)

        tt = doc.add_table(rows=1, cols=4)
        tt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["From", "To", "Type", "Status"]):
            cell = tt.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=7.5, color=WHITE)

        for idx, f in enumerate(trace_findings[:30]):
            row = tt.add_row()
            p0 = row.cells[0].paragraphs[0]
            add_run(p0, f.get("from", "—"), size=7.5, color=BLACK)
            p1 = row.cells[1].paragraphs[0]
            add_run(p1, f.get("to", "—"), size=7.5, color=BLACK)
            p2 = row.cells[2].paragraphs[0]
            add_run(p2, f.get("type", "—"), size=7.5, color=GREY_TEXT)
            p3 = row.cells[3].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p3, "MISSING", bold=True, size=7.5, color=RED_CRIT)
            style_data_row(row, idx)

        if len(trace_findings) > 30:
            p = doc.add_paragraph()
            add_run(p, f"Showing 30 of {len(trace_findings)} missing links.",
                    italic=True, size=8, color=GREY_TEXT)
    else:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        add_run(p, "\u2714  All defined traceability paths are satisfied.", bold=True,
                size=9.5, color=GREEN_OK)

    # Reference Integrity
    ri = layers.get("reference_integrity", {})
    ri_findings = ri.get("findings", [])
    if ri_findings:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        add_run(p, "6.3  Cross-Reference Integrity", bold=True, size=11, color=DARK_BLUE)
        p.paragraph_format.space_after = Pt(4)

        for f in ri_findings[:10]:
            sev = f.get("severity", "INFO")
            msg = f.get("message", "")
            p = doc.add_paragraph()
            add_run(p, f"  [{sev}] ", bold=True, size=8, color=severity_color(sev))
            add_run(p, msg[:120], size=8, color=GREY_TEXT)
            p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  7. GAP ANALYSIS
    # ═══════════════════════════════════
    add_heading(doc, "7.  Gap Analysis", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "Gaps are prioritized by risk-weighted severity using Bayesian risk scoring. "
            "Higher ASIL/risk levels receive exponential multipliers, ensuring safety-critical "
            "gaps surface first in the prioritized listing.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    gaps = gap.get("gaps", [])
    if gaps:
        gap_table = doc.add_table(rows=1, cols=5)
        gap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        gap_table.autofit = False
        gap_col_widths = [Inches(0.6), Inches(1.1), Inches(0.45), Inches(0.55), Inches(3.9)]
        for i, h in enumerate(["Priority", "Gap Category", "Count", "Risk", "Top Finding"]):
            cell = gap_table.rows[0].cells[i]
            cell.width = gap_col_widths[i]
            set_cell_shading(cell, HDR_BG)
            set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=7, color=WHITE)

        for idx, g_item in enumerate(gaps[:30]):
            priority = g_item.get("priority", "LOW")
            pri_color = RED_CRIT if priority == "HIGH" else (ACCENT if priority == "MEDIUM" else GREY_TEXT)
            key = g_item.get("key", "")
            count = g_item.get("count", 0)
            risk_w = g_item.get("risk_weighted", 0)
            top_msg = ""
            findings_list = g_item.get("findings", [])
            if findings_list:
                top_msg = findings_list[0].get("message", "")[:140]

            row = gap_table.add_row()
            for ci in range(5):
                row.cells[ci].width = gap_col_widths[ci]
                set_cell_margins(row.cells[ci], top=15, bottom=15, left=40, right=40)
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p0, priority, bold=True, size=7, color=pri_color)
            p1 = row.cells[1].paragraphs[0]
            add_run(p1, key[:45], size=6.5, color=BLACK)
            p2 = row.cells[2].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p2, str(count), size=7, color=BLACK)
            p3 = row.cells[3].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p3, str(risk_w), bold=True, size=7, color=pri_color)
            p4 = row.cells[4].paragraphs[0]
            add_run(p4, top_msg, size=6.5, color=GREY_TEXT)
            style_data_row(row, idx)

        if len(gaps) > 30:
            p = doc.add_paragraph()
            add_run(p, f"Showing 30 of {len(gaps)} gaps. Full list in Appendix A.",
                    italic=True, size=8, color=GREY_TEXT)
    else:
        p = doc.add_paragraph()
        add_run(p, "\u2714  No significant gaps identified.", bold=True, size=9.5, color=GREEN_OK)

    # ── NEW: Risk Heatmap ──
    doc.add_paragraph("")
    heatmap_path = generate_risk_heatmap(gaps, layers)
    if heatmap_path:
        _chart_files.append(heatmap_path)
        p = doc.add_paragraph()
        add_run(p, "7.1  Risk Heatmap", bold=True, size=11, color=DARK_BLUE)
        p.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        add_run(p, "The heatmap below cross-references gap categories against severity levels. "
                "Darker cells indicate higher concentrations of findings, helping prioritize "
                "remediation efforts toward the most critical compliance gaps.",
                size=9.5, color=BLACK)
        p.paragraph_format.space_after = Pt(6)
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run()
        run.add_picture(heatmap_path, width=Inches(5.0))

    doc.add_page_break()

    # ═══════════════════════════════════
    #  8. METHOD & PRACTICE AUDIT
    # ═══════════════════════════════════
    add_heading(doc, "8.  Method & Practice Audit", level=1)
    add_hr(doc)

    ma = layers.get("method_audit", {})
    ma_score = ma.get("score", 0)
    ma_req = ma.get("required_count", 0)
    ma_match = ma.get("matched_count", 0)

    p = doc.add_paragraph()
    add_run(p, "ISO 26262 specifies required and recommended methods for each ASIL level. "
            "This section audits the work products for evidence of the prescribed methods and techniques.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    add_run(p, f"For risk level ", size=9.5)
    add_run(p, f"{risk_level or 'all'}", bold=True, size=9.5)
    add_run(p, f", the standard prescribes ", size=9.5)
    add_run(p, f"{ma_req}", bold=True, size=9.5)
    add_run(p, f" methods/techniques. ", size=9.5)
    add_run(p, f"{ma_match}", bold=True, size=9.5)
    add_run(p, f" are documented (", size=9.5)
    add_run(p, f"{ma_score:.0f}%", bold=True, size=9.5, color=score_color(ma_score))
    add_run(p, f").", size=9.5)

    ma_findings = ma.get("findings", [])
    if ma_findings:
        doc.add_paragraph("")
        ma_table = doc.add_table(rows=1, cols=3)
        ma_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Severity", "Activity", "Required Methods"]):
            cell = ma_table.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=7.5, color=WHITE)

        for idx, f in enumerate(ma_findings[:20]):
            activity = f.get("activity", "")
            methods = f.get("methods", [])
            sev = f.get("severity", "WARNING")
            row = ma_table.add_row()
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p0, sev, bold=True, size=7.5, color=severity_color(sev))
            p1 = row.cells[1].paragraphs[0]
            add_run(p1, activity[:40], size=7.5, color=BLACK)
            p2 = row.cells[2].paragraphs[0]
            add_run(p2, ", ".join(methods[:5]), size=7, color=GREY_TEXT)
            style_data_row(row, idx)

    doc.add_paragraph("")

    # ═══════════════════════════════════
    #  9. RISK ASSESSMENT
    # ═══════════════════════════════════
    add_heading(doc, "9.  Risk Assessment", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "The overall risk score is computed using an ensemble of all layer results, "
            "weighted by criticality and ASIL level. Findings are categorized by severity "
            "and their cumulative risk impact is assessed.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    # Risk metrics
    total_risk = gap.get("total_risk", 0)
    add_metric_boxes(doc, [
        ("Total Risk Score", str(total_risk)),
        ("Critical Findings", str(sev_counts.get("CRITICAL", 0))),
        ("Major Findings", str(sev_counts.get("MAJOR", 0))),
        ("Warnings", str(sev_counts.get("WARNING", 0))),
    ])

    doc.add_paragraph("")

    # Concept coverage summary
    cc = layers.get("concept_coverage", {})
    cc_score_val = cc.get("score", 0)
    cc_req = cc.get("required", 0)
    cc_cov = cc.get("covered", 0)
    missing_concepts = cc.get("missing_concepts", [])

    p = doc.add_paragraph()
    add_run(p, "9.1  Concept Coverage Risk", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, f"The standard defines {cc_req} concept areas. Documentation covers {cc_cov} "
            f"({cc_score_val:.0f}%). ", size=9.5, color=BLACK)
    if missing_concepts:
        add_run(p, f"{len(missing_concepts)} concept areas are unaddressed, representing "
                "potential blind spots in the safety argument.", size=9.5, color=BLACK)

    if missing_concepts:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        add_run(p, "Missing Concept Areas:", bold=True, size=9.5, color=RED_CRIT)
        p.paragraph_format.space_after = Pt(4)

        # Display in a clean table, max 24 concepts
        display_count = min(len(missing_concepts), 24)
        cols = 3
        mc_table = doc.add_table(rows=1, cols=cols)
        mc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i in range(cols):
            cell = mc_table.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG2)
            pp = cell.paragraphs[0]
            add_run(pp, f"Concepts ({i*8+1}–{min((i+1)*8, display_count)})",
                    bold=True, size=7, color=WHITE)

        # Fill table rows (8 per column)
        rows_needed = (display_count + cols - 1) // cols
        for r_idx in range(rows_needed):
            row = mc_table.add_row()
            for c_idx in range(cols):
                flat_idx = c_idx * rows_needed + r_idx
                if flat_idx < display_count:
                    concept = missing_concepts[flat_idx]
                    pp = row.cells[c_idx].paragraphs[0]
                    add_run(pp, f"\u2022 {concept.title()}", size=7.5, color=ACCENT)
            style_data_row(row, r_idx)

        if len(missing_concepts) > display_count:
            p = doc.add_paragraph()
            add_run(p, f"  ... and {len(missing_concepts) - display_count} more. "
                    "See JSON report for complete list.",
                    italic=True, size=7.5, color=GREY_TEXT)
            p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  10. CORRECTIVE ACTION REGISTER
    # ═══════════════════════════════════
    add_heading(doc, "10.  Corrective Action Register", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "The following corrective actions are recommended based on the assessment findings. "
            "Actions are prioritized by risk impact and should be addressed in order.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    recs = summary.get("top_recommendations", gap.get("recommendations", []))
    if recs:
        ca_table = doc.add_table(rows=1, cols=4)
        ca_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["#", "Priority", "Corrective Action", "Status"]):
            cell = ca_table.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=7.5, color=WHITE)

        for idx, rec in enumerate(recs[:15]):
            if "[HIGH]" in rec:
                pri, pri_color = "HIGH", RED_CRIT
            elif "[MEDIUM]" in rec:
                pri, pri_color = "MEDIUM", ACCENT
            else:
                pri, pri_color = "LOW", GREY_TEXT

            display = rec.replace("[HIGH]", "").replace("[MEDIUM]", "").replace("[LOW]", "").strip()

            row = ca_table.add_row()
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p0, str(idx + 1), size=8, color=BLACK)
            p1 = row.cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p1, pri, bold=True, size=8, color=pri_color)
            p2 = row.cells[2].paragraphs[0]
            add_run(p2, display[:100], size=7.5, color=BLACK)
            p3 = row.cells[3].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p3, "OPEN", size=8, color=RED_CRIT)
            style_data_row(row, idx)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  11. CONFIRMATION REVIEW (ISO 26262-2:2018 §6.4.4)
    # ═══════════════════════════════════
    add_heading(doc, "11.  Confirmation Review", level=1)
    add_hr(doc)

    conf_review = report.get("confirmation_review", {})
    conf_checklist = conf_review.get("checklist", [])
    conf_summary = conf_review.get("summary", {})

    p = doc.add_paragraph()
    add_run(p, "Per ISO 26262-2:2018 §6.4.4, a confirmation review evaluates whether the work "
            "products achieve the safety objectives and comply with the applicable requirements "
            "of the standard. This section documents the confirmation review findings for the "
            "submitted work products.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    # Reviewer Independence Info
    p = doc.add_paragraph()
    add_run(p, "Reviewer Independence", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    req_independence = conf_review.get("required_independence", "I1")
    target_asil_cr = conf_review.get("target_asil", "Not specified")

    indep_table = doc.add_table(rows=1, cols=2)
    indep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Parameter", "Value"]):
        cell = indep_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG)
        pp = cell.paragraphs[0]
        add_run(pp, h, bold=True, size=8, color=WHITE)

    indep_data = [
        ("Target ASIL", str(target_asil_cr)),
        ("Required Independence Level", req_independence),
        ("I0: Same person", "Permitted for QM only"),
        ("I1: Different person, same team", "Permitted for ASIL A–B"),
        ("I2: Different department", "Required for ASIL C"),
        ("I3: Different organization (e.g. TÜV, DEKRA)", "Required for ASIL D"),
    ]
    for idx, (param, val) in enumerate(indep_data):
        row = indep_table.add_row()
        p0 = row.cells[0].paragraphs[0]
        add_run(p0, param, bold=(idx < 2), size=8, color=BLACK)
        p1 = row.cells[1].paragraphs[0]
        add_run(p1, val, bold=(idx < 2), size=8, color=DARK_BLUE if idx < 2 else GREY_TEXT)
        style_data_row(row, idx)

    doc.add_paragraph("")

    # Confirmation Checklist
    p = doc.add_paragraph()
    add_run(p, "Confirmation Review Checklist", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    if conf_checklist:
        cr_table = doc.add_table(rows=1, cols=4)
        cr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cr_table.autofit = False

        # 4-column layout: ID+Item merged, Status, Evidence (wide)
        cr_col_widths = [Inches(0.5), Inches(1.3), Inches(0.6), Inches(4.2)]
        for i, h in enumerate(["ID", "Item", "Status", "Evidence"]):
            cell = cr_table.rows[0].cells[i]
            cell.width = cr_col_widths[i]
            set_cell_shading(cell, HDR_BG)
            set_cell_margins(cell, top=30, bottom=30, left=40, right=40)
            pp = cell.paragraphs[0]
            add_run(pp, h, bold=True, size=7, color=WHITE)

        for idx, item in enumerate(conf_checklist):
            row = cr_table.add_row()
            for ci, cw in enumerate(cr_col_widths):
                row.cells[ci].width = cw
                set_cell_margins(row.cells[ci], top=15, bottom=15, left=40, right=40)
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p0, item.get("id", ""), bold=True, size=7, color=NAVY)
            p1 = row.cells[1].paragraphs[0]
            add_run(p1, item.get("item", ""), bold=True, size=7, color=BLACK)
            p2 = row.cells[2].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr_stat = item.get("status", "N/A")
            cr_stat_color = GREEN_OK if cr_stat == "Pass" else (YELLOW_WARN if cr_stat == "Partial" else RED_CRIT)
            add_run(p2, cr_stat, bold=True, size=7.5, color=cr_stat_color)
            p3 = row.cells[3].paragraphs[0]
            add_run(p3, item.get("evidence", "")[:140], size=6.5, color=GREY_TEXT)
            style_data_row(row, idx)

    doc.add_paragraph("")

    # Confirmation Review Chart
    if conf_summary:
        conf_chart_path = generate_confirmation_review_chart(conf_summary)
        if conf_chart_path:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(conf_chart_path, width=Inches(4.0))
            p.paragraph_format.space_after = Pt(8)
            _chart_files.append(conf_chart_path)

    # Confirmation Review Summary
    p = doc.add_paragraph()
    add_run(p, "Confirmation Review Summary", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    cr_result = conf_summary.get("result", "N/A")
    cr_result_color = GREEN_OK if cr_result == "PASSED" else (YELLOW_WARN if cr_result == "CONDITIONAL" else RED_CRIT)

    # Result banner
    cr_banner = doc.add_table(rows=1, cols=1)
    cr_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cr_cell = cr_banner.rows[0].cells[0]
    cr_bg = "166B34" if cr_result == "PASSED" else ("CA8A04" if cr_result == "CONDITIONAL" else "B91C1C")
    set_cell_shading(cr_cell, cr_bg)
    crp = cr_cell.paragraphs[0]
    crp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(crp, f"\n  CONFIRMATION REVIEW:  {cr_result}  \n", bold=True, size=14, color=WHITE)
    cr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    clear_borders(cr_cell)

    doc.add_paragraph("")

    cr_stats_table = doc.add_table(rows=1, cols=4)
    cr_stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Pass", "Partial", "Fail", "Pass Rate"]):
        cell = cr_stats_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG2)
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pp, h, bold=True, size=8, color=WHITE)
    row = cr_stats_table.add_row()
    for i, val in enumerate([
        str(conf_summary.get("pass", 0)),
        str(conf_summary.get("partial", 0)),
        str(conf_summary.get("fail", 0)),
        f"{conf_summary.get('pass_rate', 0):.0f}%",
    ]):
        pp = row.cells[i].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_color = GREEN_OK if i == 0 else (YELLOW_WARN if i == 1 else (RED_CRIT if i == 2 else DARK_BLUE))
        add_run(pp, val, bold=True, size=11, color=val_color)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  12. VERIFICATION REVIEW (ISO 26262-2:2018 §6.4.10 / Part 8)
    # ═══════════════════════════════════
    add_heading(doc, "12.  Verification Review", level=1)
    add_hr(doc)

    verif_review = report.get("verification_review", {})
    verif_wps = verif_review.get("work_products", [])
    verif_methods = verif_review.get("method_recommendations", [])
    verif_summary = verif_review.get("summary", {})
    verif_asil = verif_review.get("asil_key", "B")

    p = doc.add_paragraph()
    add_run(p, "Per ISO 26262-2:2018 §6.4.10 and Part 8, verification confirms that work products "
            "satisfy the requirements allocated to them. This section documents the verification "
            "methods applied to each work product and their adequacy for the target ASIL.",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    # Verification Method Recommendations Table (ISO 26262 Table notation)
    p = doc.add_paragraph()
    add_run(p, "Verification Method Recommendations by ASIL", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_run(p, "Legend:  ", bold=True, size=8, color=BLACK)
    add_run(p, "++", bold=True, size=8, color=GREEN_OK)
    add_run(p, " = Highly recommended   ", size=8, color=GREY_TEXT)
    add_run(p, "+", bold=True, size=8, color=DARK_BLUE)
    add_run(p, " = Recommended   ", size=8, color=GREY_TEXT)
    add_run(p, "o", bold=True, size=8, color=GREY_LIGHT)
    add_run(p, " = No recommendation", size=8, color=GREY_TEXT)
    p.paragraph_format.space_after = Pt(4)

    if verif_methods:
        vm_table = doc.add_table(rows=1, cols=5)
        vm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Verification Method", "ASIL A", "ASIL B", "ASIL C", "ASIL D"]):
            cell = vm_table.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG)
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(pp, h, bold=True, size=7, color=WHITE)

        for idx, vm in enumerate(verif_methods):
            row = vm_table.add_row()
            p0 = row.cells[0].paragraphs[0]
            add_run(p0, vm.get("method", ""), size=7.5, color=BLACK)
            for col_idx, asil_col in enumerate(["asil_a", "asil_b", "asil_c", "asil_d"]):
                rec_val = vm.get(asil_col, "o")
                pp = row.cells[col_idx + 1].paragraphs[0]
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rec_color = GREEN_OK if rec_val == "++" else (DARK_BLUE if rec_val == "+" else GREY_LIGHT)
                add_run(pp, rec_val, bold=True, size=8, color=rec_color)

            # Highlight the row for current ASIL
            style_data_row(row, idx)

    doc.add_paragraph("")

    # Work Product Verification Matrix
    p = doc.add_paragraph()
    add_run(p, f"Work Product Verification Status (ASIL {verif_asil})", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    if verif_wps:
        vw_table = doc.add_table(rows=1, cols=5)
        vw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        vw_table.autofit = False
        vw_col_widths = [Inches(2.2), Inches(1.0), Inches(2.0), Inches(0.7), Inches(0.7)]
        for i, h in enumerate(["Work Product", "Type", "Methods Applied", "Coverage", "Status"]):
            cell = vw_table.rows[0].cells[i]
            cell.width = vw_col_widths[i]
            set_cell_shading(cell, HDR_BG)
            set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if i >= 3 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(pp, h, bold=True, size=7, color=WHITE)

        for idx, vw in enumerate(verif_wps):
            row = vw_table.add_row()
            for ci in range(5):
                row.cells[ci].width = vw_col_widths[ci]
                set_cell_margins(row.cells[ci], top=15, bottom=15, left=40, right=40)
            p0 = row.cells[0].paragraphs[0]
            add_run(p0, vw.get("name", "")[:65], size=6.5, color=BLACK)
            p1 = row.cells[1].paragraphs[0]
            wp_type_display = vw.get("wp_type", "generic").replace("_", " ").title()[:25]
            add_run(p1, wp_type_display, size=6.5, color=GREY_TEXT)
            p2 = row.cells[2].paragraphs[0]
            methods_str = ", ".join(vw.get("methods_found", []))[:70] or "None detected"
            add_run(p2, methods_str, size=6.5, color=BLACK if methods_str != "None detected" else GREY_LIGHT)
            p3 = row.cells[3].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cov_pct = vw.get("coverage_pct", 0)
            add_run(p3, f"{cov_pct:.0f}%", bold=True, size=7.5, color=score_color(cov_pct))
            p4 = row.cells[4].paragraphs[0]
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr_stat = vw.get("status", "N/A")
            vr_color = GREEN_OK if vr_stat == "Adequate" else (YELLOW_WARN if vr_stat == "Partial" else RED_CRIT)
            add_run(p4, vr_stat, bold=True, size=7, color=vr_color)
            style_data_row(row, idx)

    doc.add_paragraph("")

    # Verification Coverage Chart
    if verif_wps:
        verif_chart_path = generate_verification_coverage_chart(verif_wps)
        if verif_chart_path:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(verif_chart_path, width=Inches(5.5))
            p.paragraph_format.space_after = Pt(8)
            _chart_files.append(verif_chart_path)

    # Verification Review Summary
    p = doc.add_paragraph()
    add_run(p, "Verification Review Summary", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    vr_result = verif_summary.get("result", "N/A")
    vr_result_color = GREEN_OK if vr_result == "PASSED" else (YELLOW_WARN if vr_result == "CONDITIONAL" else RED_CRIT)

    # Result banner
    vr_banner = doc.add_table(rows=1, cols=1)
    vr_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    vr_cell = vr_banner.rows[0].cells[0]
    vr_bg = "166B34" if vr_result == "PASSED" else ("CA8A04" if vr_result == "CONDITIONAL" else "B91C1C")
    set_cell_shading(vr_cell, vr_bg)
    vrp = vr_cell.paragraphs[0]
    vrp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(vrp, f"\n  VERIFICATION REVIEW:  {vr_result}  \n", bold=True, size=14, color=WHITE)
    vr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    clear_borders(vr_cell)

    doc.add_paragraph("")

    vr_stats_table = doc.add_table(rows=1, cols=4)
    vr_stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Adequate", "Partial", "Insufficient", "Adequacy Rate"]):
        cell = vr_stats_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG2)
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pp, h, bold=True, size=8, color=WHITE)
    row = vr_stats_table.add_row()
    for i, val in enumerate([
        str(verif_summary.get("adequate", 0)),
        str(verif_summary.get("partial", 0)),
        str(verif_summary.get("insufficient", 0)),
        f"{verif_summary.get('adequacy_rate', 0):.0f}%",
    ]):
        pp = row.cells[i].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_color = GREEN_OK if i == 0 else (YELLOW_WARN if i == 1 else (RED_CRIT if i == 2 else DARK_BLUE))
        add_run(pp, val, bold=True, size=11, color=val_color)

    # Missing methods detail (only if there are insufficiencies)
    missing_methods_wps = [vw for vw in verif_wps if vw.get("methods_missing")]
    if missing_methods_wps:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        add_run(p, "Missing Verification Methods", bold=True, size=10, color=ACCENT)
        p.paragraph_format.space_after = Pt(4)

        mm_table = doc.add_table(rows=1, cols=3)
        mm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Work Product", "Missing Methods", "Rec. Level"]):
            cell = mm_table.rows[0].cells[i]
            set_cell_shading(cell, HDR_BG2)
            pp = cell.paragraphs[0]
            add_run(pp, h, bold=True, size=7.5, color=WHITE)

        for idx, vw in enumerate(missing_methods_wps[:20]):
            for method in vw.get("methods_missing", [])[:3]:
                row = mm_table.add_row()
                p0 = row.cells[0].paragraphs[0]
                add_run(p0, vw.get("name", "")[:30], size=7, color=BLACK)
                p1 = row.cells[1].paragraphs[0]
                add_run(p1, method, size=7, color=RED_CRIT)
                p2 = row.cells[2].paragraphs[0]
                # Look up recommendation level
                rec_lvl = "o"
                for md in vw.get("method_details", []):
                    if md.get("method") == method:
                        rec_lvl = md.get("recommendation", "o")
                        break
                rec_color = GREEN_OK if rec_lvl == "++" else (DARK_BLUE if rec_lvl == "+" else GREY_LIGHT)
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p2, rec_lvl, bold=True, size=8, color=rec_color)
                style_data_row(row, idx)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  13. ASSESSMENT DECISION
    # ═══════════════════════════════════
    add_heading(doc, "13.  Assessment Decision", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "Based on the comprehensive analysis performed across all 8 algorithm layers, "
            "including requirements coverage, content alignment, semantic matching, concept "
            "coverage, reference integrity, method audit, traceability chain analysis, and "
            "risk-weighted gap analysis, the following assessment decision is rendered:",
            size=9.5, color=BLACK)
    p.paragraph_format.space_after = Pt(12)

    # Verdict banner
    v_table2 = doc.add_table(rows=1, cols=1)
    v_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    v_cell2 = v_table2.rows[0].cells[0]
    set_cell_shading(v_cell2, v_hex)
    vp2 = v_cell2.paragraphs[0]
    vp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(vp2, f"\n  ASSESSMENT VERDICT:  {verdict}  \n", bold=True, size=18, color=WHITE)
    v_cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(v_cell2, top="0 none FFFFFF", bottom="0 none FFFFFF",
                    left="0 none FFFFFF", right="0 none FFFFFF")

    doc.add_paragraph("")

    p = doc.add_paragraph()
    add_run(p, decision.get("description", ""), size=10, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    if decision.get("conditions"):
        p = doc.add_paragraph()
        add_run(p, "Conditions for Acceptance:", bold=True, size=10, color=ACCENT)
        p.paragraph_format.space_after = Pt(4)
        for cond in decision["conditions"]:
            p = doc.add_paragraph()
            add_run(p, f"    \u2022  {cond}", size=9.5, color=BLACK)
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Signature block
    p = doc.add_paragraph()
    add_run(p, "Signatures", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(8)

    sig_table = doc.add_table(rows=4, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Role", "Name", "Signature / Date"]):
        cell = sig_table.rows[0].cells[i]
        set_cell_shading(cell, HDR_BG2)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=8, color=WHITE)

    sig_roles = ["Assessor", "Safety Manager", "Project Manager"]
    for idx, role in enumerate(sig_roles):
        row = sig_table.rows[idx + 1]
        p0 = row.cells[0].paragraphs[0]
        add_run(p0, role, bold=True, size=9, color=BLACK)
        p1 = row.cells[1].paragraphs[0]
        add_run(p1, "________________________________", size=9, color=GREY_LIGHT)
        p2 = row.cells[2].paragraphs[0]
        add_run(p2, "________________________________", size=9, color=GREY_LIGHT)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  APPENDIX A: DETAILED FINDINGS
    # ═══════════════════════════════════
    add_heading(doc, "Appendix A:  Detailed Findings", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "Complete listing of all findings from all algorithm layers, sorted by severity.",
            size=9, color=BLACK)
    p.paragraph_format.space_after = Pt(6)

    all_findings = []
    for layer_name, layer_data in layers.items():
        for f in layer_data.get("findings", []):
            f_copy = dict(f)
            f_copy["_layer"] = layer_name
            all_findings.append(f_copy)

    sev_order = {"CRITICAL": 0, "MAJOR": 1, "WARNING": 2, "INFO": 3}
    all_findings.sort(key=lambda f: sev_order.get(f.get("severity", "INFO"), 4))

    if all_findings:
        app_table = doc.add_table(rows=1, cols=4)
        app_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        app_table.autofit = False

        # Column widths: Severity(0.7), Layer(1.0), Type(1.2), Details(3.7)
        app_col_widths = [Inches(0.7), Inches(1.0), Inches(1.2), Inches(3.7)]
        for i, h in enumerate(["Severity", "Layer", "Type", "Details"]):
            cell = app_table.rows[0].cells[i]
            cell.width = app_col_widths[i]
            set_cell_shading(cell, HDR_BG)
            set_cell_margins(cell, top=30, bottom=30, left=40, right=40)
            p = cell.paragraphs[0]
            add_run(p, h, bold=True, size=7, color=WHITE)

        for idx, f in enumerate(all_findings[:100]):
            sev = f.get("severity", "INFO")
            layer = f.get("_layer", "").replace("_", " ").title()[:20]
            ftype = f.get("type", "")[:24]
            msg = f.get("message", "")[:160]  # Allow longer details

            row = app_table.add_row()
            for ci, cw in enumerate(app_col_widths):
                row.cells[ci].width = cw
                set_cell_margins(row.cells[ci], top=15, bottom=15, left=40, right=40)
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p0, sev, bold=True, size=6.5, color=severity_color(sev))
            p1 = row.cells[1].paragraphs[0]
            add_run(p1, layer, size=6.5, color=BLACK)
            p2 = row.cells[2].paragraphs[0]
            add_run(p2, ftype, size=6.5, color=BLACK)
            p3 = row.cells[3].paragraphs[0]
            add_run(p3, msg, size=6.5, color=GREY_TEXT)
            style_data_row(row, idx)

        if len(all_findings) > 100:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            add_run(p, f"Showing 100 of {len(all_findings)} findings. "
                    "Complete data available in JSON report.",
                    italic=True, size=7.5, color=GREY_TEXT)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  APPENDIX B: ALGORITHM LAYER DETAILS
    # ═══════════════════════════════════
    add_heading(doc, "Appendix B:  Algorithm Layer Details", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "Detailed results from each of the 8 algorithm layers.",
            size=9, color=BLACK)
    p.paragraph_format.space_after = Pt(8)

    layer_order = ["node_coverage", "content_alignment", "semantic_matching",
                   "concept_coverage", "reference_integrity", "method_audit",
                   "traceability_chain"]

    for lkey in layer_order:
        ldata = layers.get(lkey, {})
        if not ldata:
            continue
        lname = layer_labels.get(lkey, lkey.replace("_", " ").title())
        lscore = ldata.get("score", 0)

        p = doc.add_paragraph()
        add_run(p, f"{lname}", bold=True, size=10, color=DARK_BLUE)
        add_run(p, f"  —  {lscore:.0f}%", bold=True, size=10, color=score_color(lscore))
        p.paragraph_format.space_after = Pt(2)

        # Key stats
        stats = []
        for k, v in ldata.items():
            if k in ("score", "layer", "findings"):
                continue
            if isinstance(v, (int, float, str)) and v != "":
                stats.append(f"{k.replace('_', ' ').title()}: {v}")
        if stats:
            p = doc.add_paragraph()
            add_run(p, "  " + "  |  ".join(stats[:6]), size=8, color=GREY_TEXT)
            p.paragraph_format.space_after = Pt(2)

        n_findings = len(ldata.get("findings", []))
        p = doc.add_paragraph()
        add_run(p, f"  Findings: {n_findings}", size=8, color=GREY_TEXT)
        p.paragraph_format.space_after = Pt(6)

        add_thin_hr(doc)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  APPENDIX C: GLOSSARY & REFERENCES
    # ═══════════════════════════════════
    add_heading(doc, "Appendix C:  Glossary & References", level=1)
    add_hr(doc)

    p = doc.add_paragraph()
    add_run(p, "References", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    references = [
        "ISO 26262:2018 — Road vehicles — Functional safety (Parts 1–12)",
        "ISO/SAE 21434:2021 — Road vehicles — Cybersecurity engineering",
        "ISO 21448:2022 — Road vehicles — Safety of the intended functionality (SOTIF)",
        "IEC 61508:2010 — Functional safety of electrical/electronic/programmable electronic safety-related systems",
        "Automotive SPICE (PAM 3.1) — Process Assessment Model",
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        add_run(p, f"  [{i}]  ", bold=True, size=9, color=NAVY)
        add_run(p, ref, size=9, color=BLACK)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    add_run(p, "Glossary", bold=True, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    glossary = [
        ("ASIL", "Automotive Safety Integrity Level — classification scheme (QM, A, B, C, D)"),
        ("Compliance Score", "Weighted aggregate of all algorithm layer scores (0–100%)"),
        ("Diagnostic Coverage", "Fraction of faults detectable by safety mechanisms"),
        ("FMEA", "Failure Mode and Effects Analysis"),
        ("FTA", "Fault Tree Analysis — top-down deductive failure analysis"),
        ("HARA", "Hazard Analysis and Risk Assessment — systematic process to identify hazards"),
        ("PMHF", "Probabilistic Metric for Random Hardware Failures"),
        ("Safety Case", "Structured argument demonstrating system safety"),
        ("Safety Goal", "Top-level safety requirement derived from HARA"),
        ("Traceability", "Ability to trace requirements through all development phases"),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph()
        add_run(p, f"  {term}: ", bold=True, size=9, color=NAVY)
        add_run(p, defn, size=9, color=GREY_TEXT)
        p.paragraph_format.space_after = Pt(2)

    # ── Footer ──
    doc.add_paragraph("")
    add_hr(doc, color_hex="D1D5DB", weight="4")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, f"Lion of Functional Safety Engine\u2122 — Compliance Assessment Report\n"
            f"Document {doc_id}  |  Generated {check_date}\n"
            f"Page numbers and section cross-references are for electronic viewing.",
            size=7, color=GREY_LIGHT)

    # ── Save ──
    doc.save(output_path)

    # ── Cleanup temp chart files ──
    for cf in _chart_files:
        try:
            os.unlink(cf)
        except OSError:
            pass

    print(f"  [OK] DOCX report generated: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python report_generator_docx.py <compliance_report.json> [output.docx]")
        sys.exit(1)

    report_path = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else report_path.replace(".json", ".docx")

    with open(report_path, 'r') as f:
        report = json.load(f)

    generate_compliance_docx(report, output)
